"""
Generate comprehensive real estate document sets for 10 projects.

Project 01 — Sunrise Heights — gets ALL 8 document types:
  brochure.pdf, price_list.xlsx, floor_plans.pdf, amenities_guide.pdf,
  payment_plan.pdf, legal_rera.pdf, faq.docx, possession_timeline.xlsx

Projects 02–10 — get 4 document types each:
  brochure.pdf, price_list.xlsx, amenities_guide.pdf, faq.docx

Run from repo root:
  python scripts/generate_project_docs.py
"""

import os
from fpdf import FPDF
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = "./real_estate_documents"


# =============================================================================
# PROJECT DATA
# =============================================================================

PROJECTS = [
    {
        "folder": "01_Sunrise_Heights",
        "name": "Sunrise Heights Premium Apartments",
        "short": "Sunrise Heights",
        "developer": "Prestige Group",
        "location": "Whitefield, Bangalore - 560066",
        "total_units": 520, "towers": 4, "floors": 22, "acres": 5.5,
        "rera": "PRM/KA/RERA/1251/309/PR/210523/004789",
        "construction_start": "January 2023", "possession": "December 2026",
        "prop_type": "Apartment",
        "nearby": [
            "Whitefield Metro Station - 1.2 km",
            "ITPL Tech Park - 2.5 km",
            "Phoenix Marketcity Mall - 3.0 km",
            "Manipal Hospital - 1.8 km",
            "Inventure Academy School - 1.5 km",
            "Whitefield Railway Station - 2.0 km",
        ],
        "units": [
            {"type": "2BHK Standard",  "area": 1050, "price_sqft": 7500, "count": 120, "bath": 2},
            {"type": "2BHK Premium",   "area": 1200, "price_sqft": 8000, "count": 100, "bath": 2},
            {"type": "3BHK Standard",  "area": 1550, "price_sqft": 7800, "count": 160, "bath": 3},
            {"type": "3BHK Premium",   "area": 1800, "price_sqft": 8200, "count": 80,  "bath": 3},
            {"type": "3BHK Duplex",    "area": 2200, "price_sqft": 8500, "count": 40,  "bath": 4},
            {"type": "4BHK Penthouse", "area": 3200, "price_sqft": 9500, "count": 20,  "bath": 5},
        ],
        "amenities": [
            "Olympic-size swimming pool (50 m)",
            "Fully-equipped gymnasium (5,000 sq.ft.)",
            "Clubhouse (15,000 sq.ft.)",
            "Indoor badminton courts (2 nos.)",
            "Squash court",
            "Yoga and meditation deck",
            "Children's play area with soft-fall surface",
            "Senior citizen sit-out area",
            "Jogging / cycling track (400 m)",
            "Landscaped gardens (2 acres)",
            "Multipurpose hall (capacity 300 pax)",
            "Mini theatre (50 seats)",
            "Library and reading lounge",
            "Co-working space (30 workstations)",
            "Rooftop sky lounge",
            "BBQ and outdoor dining area",
            "Cricket pitch and net practice area",
            "Basketball court",
            "Tennis court (2 nos.)",
            "Table tennis room",
            "24x7 CCTV surveillance (200+ cameras)",
            "3-tier security — boom barrier, biometric, guard patrolling",
            "EV charging stations (50 nos.)",
            "Solar-powered common areas",
            "Rainwater harvesting (2 lakh litre capacity)",
            "Sewage treatment plant (STP)",
            "Organic waste converter",
            "Underground parking — 1,040 slots (2 per unit)",
            "Visitor parking — 80 slots",
            "24x7 concierge and housekeeping",
        ],
        "full_set": True,
    },
    {
        "folder": "02_Green_Valley_Villas",
        "name": "Green Valley Luxury Villas",
        "short": "Green Valley",
        "developer": "Sobha Developers",
        "location": "Devanahalli, Bangalore - 562110",
        "total_units": 120, "towers": 0, "floors": 2, "acres": 18.0,
        "rera": "PRM/KA/RERA/1251/446/PR/180423/005612",
        "construction_start": "March 2023", "possession": "June 2026",
        "prop_type": "Villa",
        "nearby": [
            "Kempegowda International Airport - 8 km",
            "Devanahalli Business Park - 3 km",
            "BIAL Aerospace SEZ - 5 km",
            "Devanahalli Town - 2.5 km",
            "Columbia Asia Hospital - 6 km",
        ],
        "units": [
            {"type": "3BHK Villa",     "area": 2200, "price_sqft": 6800, "count": 50, "bath": 3},
            {"type": "4BHK Villa",     "area": 3000, "price_sqft": 7200, "count": 45, "bath": 4},
            {"type": "5BHK Grand Villa","area": 4200, "price_sqft": 8000, "count": 25, "bath": 5},
        ],
        "amenities": [
            "Private swimming pool per villa (optional)",
            "Community clubhouse (8,000 sq.ft.)",
            "Infinity pool with deck",
            "Spa and wellness centre",
            "Children's play zone (1 acre)",
            "Nature trails and walking paths (1.5 km)",
            "Basketball and tennis court",
            "Organic farming plots for residents",
            "24x7 gated security with CCTV",
            "Underground power cables",
            "Rain water harvesting",
            "Sewage treatment plant",
        ],
        "full_set": False,
    },
    {
        "folder": "03_Metro_Edge_Compact_Homes",
        "name": "Metro Edge Compact Homes",
        "short": "Metro Edge",
        "developer": "Puravankara Limited",
        "location": "Thanisandra Main Road, Bangalore - 560077",
        "total_units": 840, "towers": 6, "floors": 18, "acres": 4.2,
        "rera": "PRM/KA/RERA/1251/388/PR/290323/006123",
        "construction_start": "June 2023", "possession": "March 2027",
        "prop_type": "Apartment",
        "nearby": [
            "Nagawara Metro Station - 0.8 km",
            "Manyata Tech Park - 3.5 km",
            "Hebbal Flyover - 4 km",
            "Columbia Asia Hospital Hebbal - 3 km",
            "Elements Mall - 2 km",
        ],
        "units": [
            {"type": "1BHK",          "area": 600,  "price_sqft": 6500, "count": 280, "bath": 1},
            {"type": "2BHK Compact",  "area": 900,  "price_sqft": 6800, "count": 360, "bath": 2},
            {"type": "2BHK Standard", "area": 1100, "price_sqft": 7000, "count": 200, "bath": 2},
        ],
        "amenities": [
            "Swimming pool and kids pool",
            "Gymnasium",
            "Multipurpose sports court",
            "Children's play area",
            "Landscaped podium garden",
            "Co-working lounge",
            "24x7 security",
            "Power backup for common areas",
            "EV charging points (20 nos.)",
            "Rooftop terrace garden",
        ],
        "full_set": False,
    },
    {
        "folder": "04_Royal_Orchid_Residency",
        "name": "Royal Orchid Residency",
        "short": "Royal Orchid",
        "developer": "Brigade Group",
        "location": "Sarjapur Road, Bangalore - 560035",
        "total_units": 380, "towers": 3, "floors": 25, "acres": 3.8,
        "rera": "PRM/KA/RERA/1251/512/PR/140223/007890",
        "construction_start": "February 2023", "possession": "September 2026",
        "prop_type": "Apartment",
        "nearby": [
            "Carmelaram Railway Station - 1.5 km",
            "Embassy Tech Village - 4 km",
            "RMZ Ecospace - 5 km",
            "Iblur Junction - 2 km",
            "Vibgyor High School - 1.8 km",
        ],
        "units": [
            {"type": "2BHK",          "area": 1150, "price_sqft": 8200, "count": 150, "bath": 2},
            {"type": "3BHK Standard", "area": 1700, "price_sqft": 8500, "count": 160, "bath": 3},
            {"type": "3BHK Large",    "area": 2100, "price_sqft": 9000, "count": 70,  "bath": 3},
        ],
        "amenities": [
            "Temperature-controlled swimming pool",
            "Fully equipped gym",
            "Aerobics and dance studio",
            "Indoor games room",
            "Banquet hall (200 pax)",
            "Mini theatre",
            "Jogging track",
            "Yoga lawn",
            "Pet-friendly park",
            "24x7 concierge",
            "Solar water heating",
            "Rainwater harvesting",
        ],
        "full_set": False,
    },
    {
        "folder": "05_Eco_Habitat_Township",
        "name": "Eco Habitat Integrated Township",
        "short": "Eco Habitat",
        "developer": "Godrej Properties",
        "location": "Doddaballapur Road, Bangalore - 561203",
        "total_units": 1200, "towers": 8, "floors": 15, "acres": 42.0,
        "rera": "PRM/KA/RERA/1251/603/PR/010623/009012",
        "construction_start": "September 2022", "possession": "December 2025",
        "prop_type": "Township",
        "nearby": [
            "Doddaballapur Town Centre - 2 km",
            "Hardware City SEZ - 5 km",
            "BIAL Airport - 22 km",
            "Bangalore City Centre - 35 km",
            "Industrial area Doddaballapur - 3 km",
        ],
        "units": [
            {"type": "1BHK Economy",  "area": 500,  "price_sqft": 4800, "count": 300, "bath": 1},
            {"type": "2BHK Standard", "area": 850,  "price_sqft": 5200, "count": 500, "bath": 2},
            {"type": "2BHK Premium",  "area": 1050, "price_sqft": 5600, "count": 300, "bath": 2},
            {"type": "3BHK",          "area": 1400, "price_sqft": 6000, "count": 100, "bath": 3},
        ],
        "amenities": [
            "3 swimming pools across phases",
            "Gymnasium (3 nos.)",
            "Schools within campus (K-12)",
            "Retail high street (2 km stretch)",
            "Multi-speciality clinic",
            "Jogging tracks (5 km combined)",
            "Community parks (10 acres)",
            "Shopping supermarket",
            "Food court",
            "EV charging hub",
            "100% solar common areas",
            "Zero waste management",
        ],
        "full_set": False,
    },
    {
        "folder": "06_Lakeshore_Towers",
        "name": "Lakeshore Towers Luxury Residences",
        "short": "Lakeshore Towers",
        "developer": "Embassy Group",
        "location": "Hebbal, Bangalore - 560024",
        "total_units": 240, "towers": 2, "floors": 30, "acres": 2.8,
        "rera": "PRM/KA/RERA/1251/718/PR/220423/010234",
        "construction_start": "April 2023", "possession": "June 2027",
        "prop_type": "Luxury Apartment",
        "nearby": [
            "Hebbal Lake - 0.3 km (lake view units available)",
            "Manyata Tech Park - 2 km",
            "Esteem Mall - 1.5 km",
            "Bangalore International Airport - 18 km",
            "Yashwantpur Junction - 8 km",
        ],
        "units": [
            {"type": "3BHK Lake View",   "area": 2000, "price_sqft": 11000, "count": 80,  "bath": 3},
            {"type": "3BHK City View",   "area": 1900, "price_sqft": 10000, "count": 90,  "bath": 3},
            {"type": "4BHK Luxury",      "area": 2800, "price_sqft": 12500, "count": 50,  "bath": 4},
            {"type": "4BHK Penthouse",   "area": 4500, "price_sqft": 15000, "count": 20,  "bath": 5},
        ],
        "amenities": [
            "Lakeside infinity pool",
            "Private cabanas and sun deck",
            "Signature spa and wellness centre",
            "Fine dining restaurant (residents only)",
            "Wine cellar and tasting lounge",
            "Private cinema hall (30 seats)",
            "Concierge and butler services",
            "Helipad (on approval)",
            "Business centre (conference rooms)",
            "Golf simulator",
            "Squash and tennis courts",
            "Smart home automation",
            "5-tier security",
            "Branded luxury lobbies",
        ],
        "full_set": False,
    },
    {
        "folder": "07_Heritage_Grand_Senior_Living",
        "name": "Heritage Grand Senior Living Community",
        "short": "Heritage Grand",
        "developer": "Ashiana Housing",
        "location": "Yelahanka, Bangalore - 560064",
        "total_units": 200, "towers": 3, "floors": 5, "acres": 6.0,
        "rera": "PRM/KA/RERA/1251/829/PR/050523/011456",
        "construction_start": "May 2023", "possession": "March 2026",
        "prop_type": "Senior Living Apartment",
        "nearby": [
            "Yelahanka New Town - 1 km",
            "BIAL Airport - 15 km",
            "Narayana Multispeciality Hospital - 2 km",
            "Yelahanka Lake - 1.5 km",
            "CRPF Campus - 1 km",
        ],
        "units": [
            {"type": "1BHK Senior Suite", "area": 700,  "price_sqft": 6500, "count": 80,  "bath": 1},
            {"type": "2BHK Senior",       "area": 1050, "price_sqft": 7000, "count": 90,  "bath": 2},
            {"type": "2BHK Premium",      "area": 1250, "price_sqft": 7500, "count": 30,  "bath": 2},
        ],
        "amenities": [
            "24x7 on-site medical clinic with visiting doctors",
            "Dedicated ambulance service",
            "Therapeutic swimming pool (heated)",
            "Physiotherapy and rehabilitation centre",
            "Memory care programme",
            "Yoga and light exercise classes",
            "Library and reading room",
            "Chess and board games lounge",
            "Cultural event hall",
            "Vegetable and flower garden (resident-maintained)",
            "All-day dining restaurant with dietary options",
            "Housekeeping included in maintenance",
            "Emergency call system in all units",
            "CCTV surveillance",
            "Geriatric-friendly design (ramps, wide corridors, grab bars)",
        ],
        "full_set": False,
    },
    {
        "folder": "08_Skyline_Commercial_Plaza",
        "name": "Skyline Commercial Plaza",
        "short": "Skyline Plaza",
        "developer": "Salarpuria Sattva Group",
        "location": "Outer Ring Road, Marathahalli, Bangalore - 560037",
        "total_units": 320, "towers": 1, "floors": 18, "acres": 2.0,
        "rera": "PRM/KA/RERA/1251/934/PR/120423/012678",
        "construction_start": "April 2023", "possession": "January 2026",
        "prop_type": "Commercial Office Space",
        "nearby": [
            "Marathahalli Bridge - 0.5 km",
            "Bagmane Tech Park - 2 km",
            "RMZ Infinity - 3 km",
            "ITPL Whitefield - 5 km",
            "Outer Ring Road Metro (proposed) - 0.8 km",
        ],
        "units": [
            {"type": "Small Office (300 sq.ft.)",  "area": 300,  "price_sqft": 12000, "count": 100, "bath": 1},
            {"type": "Medium Office (600 sq.ft.)", "area": 600,  "price_sqft": 11500, "count": 120, "bath": 2},
            {"type": "Large Office (1200 sq.ft.)", "area": 1200, "price_sqft": 11000, "count": 70,  "bath": 2},
            {"type": "Floor Plate (5000 sq.ft.)",  "area": 5000, "price_sqft": 10000, "count": 30,  "bath": 4},
        ],
        "amenities": [
            "Grade-A office specifications",
            "Centrally air-conditioned (VRV system)",
            "High-speed elevators (8 nos.)",
            "Power backup (100% DG)",
            "24x7 access control",
            "EV charging stations (40 nos.)",
            "Multi-level car parking (640 slots)",
            "Food court (ground floor, 8 outlets)",
            "ATM and banking zone",
            "Conference centre (200 pax)",
            "Crèche facility",
            "LEED Gold certified building",
            "Rainwater harvesting",
            "100 Mbps internet backbone",
        ],
        "full_set": False,
    },
    {
        "folder": "09_Palm_Grove_Plots",
        "name": "Palm Grove Residential Plots",
        "short": "Palm Grove Plots",
        "developer": "Mahindra Lifespaces",
        "location": "Kanakapura Road, Bangalore - 560062",
        "total_units": 450, "towers": 0, "floors": 0, "acres": 55.0,
        "rera": "PRM/KA/RERA/1251/1047/PR/300323/013890",
        "construction_start": "N/A (ready-to-register plots)", "possession": "Immediate Registration",
        "prop_type": "Residential Plot",
        "nearby": [
            "Art of Living Ashram - 3 km",
            "NICE Road (Peripheral Ring Road) - 2 km",
            "Kanakapura Town - 5 km",
            "Bangalore City Centre - 22 km",
            "Talaghattapura Metro Station - 8 km",
        ],
        "units": [
            {"type": "30x40 Plot (1,200 sq.ft.)", "area": 1200, "price_sqft": 3200, "count": 180, "bath": 0},
            {"type": "30x50 Plot (1,500 sq.ft.)", "area": 1500, "price_sqft": 3000, "count": 150, "bath": 0},
            {"type": "40x60 Plot (2,400 sq.ft.)", "area": 2400, "price_sqft": 2800, "count": 80,  "bath": 0},
            {"type": "60x80 Corner Plot (4,800 sq.ft.)", "area": 4800, "price_sqft": 2600, "count": 40, "bath": 0},
        ],
        "amenities": [
            "Fully formed BDA-approved roads (40 ft and 30 ft)",
            "Underground drainage system",
            "24x7 potable water supply",
            "Underground electrical cables",
            "Street lighting throughout layout",
            "Landscaped parks (3 nos. across layout)",
            "Children's play equipment",
            "Tree-lined avenues",
            "Compound wall with gated entrance",
            "24x7 security",
            "CCTV at entry and exit points",
            "Clubhouse (2,500 sq.ft.) for plot owners",
        ],
        "full_set": False,
    },
    {
        "folder": "10_Prestige_Park_Grove",
        "name": "Prestige Park Grove",
        "short": "Park Grove",
        "developer": "Prestige Group",
        "location": "Whitefield - Budigere Road, Bangalore - 560049",
        "total_units": 3000, "towers": 12, "floors": 20, "acres": 80.0,
        "rera": "PRM/KA/RERA/1251/1158/PR/150223/015012",
        "construction_start": "February 2023", "possession": "December 2027",
        "prop_type": "Integrated Township",
        "nearby": [
            "Budigere Cross - 1 km",
            "Whitefield - 8 km",
            "ITPL Tech Park - 10 km",
            "Hoskote Industrial Area - 6 km",
            "Old Madras Road - 3 km",
        ],
        "units": [
            {"type": "1BHK",          "area": 550,  "price_sqft": 5800, "count": 600,  "bath": 1},
            {"type": "2BHK Standard", "area": 900,  "price_sqft": 6200, "count": 1000, "bath": 2},
            {"type": "2BHK Premium",  "area": 1100, "price_sqft": 6500, "count": 800,  "bath": 2},
            {"type": "3BHK",          "area": 1450, "price_sqft": 7000, "count": 500,  "bath": 3},
            {"type": "4BHK Penthouse","area": 2800, "price_sqft": 8500, "count": 100,  "bath": 4},
        ],
        "amenities": [
            "5 swimming pools (Olympic + leisure + kids)",
            "Sports complex (cricket, football, basketball, tennis)",
            "Gymnasium (4 nos.)",
            "International school within campus",
            "Multi-speciality hospital tie-up",
            "Shopping mall (2 lakh sq.ft.)",
            "Food courts (3 nos.)",
            "Amphitheatre (capacity 500 pax)",
            "Jogging and cycling tracks (8 km)",
            "Parks and open spaces (25 acres)",
            "EV charging hub (100 stations)",
            "Solar plant (2 MW)",
            "STP and water recycling",
            "24x7 multi-tier security",
        ],
        "full_set": False,
    },
]


# =============================================================================
# PDF HELPER
# =============================================================================

def _safe(text: str) -> str:
    """Replace non-latin-1 characters with ASCII equivalents for Helvetica."""
    return (text
        .replace("\u2014", "--")   # em dash
        .replace("\u2013", "-")    # en dash
        .replace("\u2018", "'").replace("\u2019", "'")  # smart quotes
        .replace("\u201c", '"').replace("\u201d", '"')
        .replace("\u2022", "-")    # bullet
        .replace("\u00d7", "x")    # multiplication sign
        .encode("latin-1", errors="replace").decode("latin-1")
    )


def new_pdf():
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    return pdf


def pdf_h1(pdf, text):
    pdf.set_font("Helvetica", "B", 18)
    pdf.cell(0, 12, _safe(text), new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(2)


def pdf_h2(pdf, text):
    pdf.set_font("Helvetica", "B", 13)
    pdf.set_fill_color(240, 240, 240)
    pdf.cell(0, 9, _safe(f"  {text}"), new_x="LMARGIN", new_y="NEXT", fill=True)
    pdf.ln(2)


def _w(pdf):
    """Usable text width."""
    return pdf.w - pdf.l_margin - pdf.r_margin


def pdf_body(pdf, text):
    pdf.set_font("Helvetica", "", 10)
    pdf.set_x(pdf.l_margin)
    pdf.multi_cell(_w(pdf), 6, _safe(text))
    pdf.ln(2)


def pdf_bullet(pdf, items):
    pdf.set_font("Helvetica", "", 10)
    for item in items:
        pdf.set_x(pdf.l_margin)
        pdf.multi_cell(_w(pdf), 6, _safe(f"  -  {item}"))
    pdf.ln(2)


def pdf_kv(pdf, label, value):
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(55, 7, _safe(label + ":"), new_x="RIGHT", new_y="LAST")
    pdf.set_font("Helvetica", "", 10)
    pdf.cell(0, 7, _safe(str(value)), new_x="LMARGIN", new_y="NEXT")


def lakhs(amount):
    l = amount / 100000
    if l >= 100:
        return f"Rs. {l/100:.2f} Crores"
    return f"Rs. {l:.2f} Lakhs"


# =============================================================================
# DOCUMENT 1 — PROJECT BROCHURE (PDF)  — ALL PROJECTS
# =============================================================================

def create_brochure(p, out_dir):
    pdf = new_pdf()
    pdf.add_page()

    # Cover
    pdf_h1(pdf, p["name"])
    pdf.set_font("Helvetica", "I", 11)
    pdf.cell(0, 8, p["location"], new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.cell(0, 7, f"By {p['developer']}", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(6)

    # Overview
    pdf_h2(pdf, "Project Overview")
    total_area = sum(u["area"] * u["count"] for u in p["units"])
    body = (
        f"{p['name']} is a {p['prop_type'].lower()} project spread across "
        f"{p['acres']} acres in {p['location']}, developed by {p['developer']}. "
    )
    if p["towers"] > 0:
        body += (
            f"The project comprises {p['towers']} towers with {p['floors']} floors each, "
            f"housing {p['total_units']} units. "
        )
    else:
        body += f"The project offers {p['total_units']} {p['prop_type'].lower()} units totalling {total_area:,} sq.ft. "

    body += (
        f"Construction commenced in {p['construction_start']} and possession is expected by {p['possession']}. "
        f"RERA Registration No.: {p['rera']}."
    )
    pdf_body(pdf, body)

    # Key Details
    pdf_h2(pdf, "Key Project Details")
    pdf_kv(pdf, "Developer", p["developer"])
    pdf_kv(pdf, "Location", p["location"])
    pdf_kv(pdf, "Property Type", p["prop_type"])
    pdf_kv(pdf, "Total Units", str(p["total_units"]))
    pdf_kv(pdf, "Land Area", f"{p['acres']} acres")
    if p["towers"] > 0:
        pdf_kv(pdf, "Towers / Floors", f"{p['towers']} towers x {p['floors']} floors")
    pdf_kv(pdf, "Construction Start", p["construction_start"])
    pdf_kv(pdf, "Possession", p["possession"])
    pdf_kv(pdf, "RERA No.", p["rera"])
    pdf.ln(3)

    # Unit Types
    pdf.add_page()
    pdf_h2(pdf, "Available Unit Configurations")
    for u in p["units"]:
        base = u["area"] * u["price_sqft"]
        plc = int(base * 0.05)
        total = base + plc + 300000
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 7, f"  {u['type']}", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(0, 6, f"     Super Built-up Area: {u['area']:,} sq.ft.  |  "
                       f"Bathrooms: {u['bath']}  |  Available Units: {u['count']}", new_x="LMARGIN", new_y="NEXT")
        pdf.cell(0, 6, f"     Base Price: Rs. {u['price_sqft']:,}/sq.ft.  |  "
                       f"Total (approx.): {lakhs(total)}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

    # Location Advantages
    pdf_h2(pdf, "Location Advantages")
    pdf_bullet(pdf, p["nearby"])

    # Why Choose
    pdf_h2(pdf, "Why Choose " + p["short"] + "?")
    reasons = [
        f"Premium {p['prop_type'].lower()} by {p['developer']}, a trusted brand with 30+ years of excellence",
        "RERA-registered project ensuring transparency and buyer protection",
        f"Strategic location at {p['location']} with excellent connectivity",
        "World-class amenities for a complete lifestyle experience",
        "Modern architecture with Vastu-compliant layouts",
        "24x7 security for complete peace of mind",
        "Green-certified construction with sustainable features",
    ]
    pdf_bullet(pdf, reasons)

    path = os.path.join(out_dir, "01_brochure.pdf")
    pdf.output(path)
    print(f"  ✓ {path}")


# =============================================================================
# DOCUMENT 2 — PRICE LIST (XLSX) — ALL PROJECTS
# =============================================================================

def create_price_list(p, out_dir):
    wb = Workbook()
    ws = wb.active
    ws.title = "Price List"

    # Title
    ws.merge_cells("A1:I1")
    ws["A1"] = f"{p['name']} — Price List"
    ws["A1"].font = Font(bold=True, size=14)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 28

    ws.merge_cells("A2:I2")
    ws["A2"] = f"Location: {p['location']}  |  Developer: {p['developer']}  |  RERA: {p['rera']}"
    ws["A2"].font = Font(italic=True, size=10)
    ws["A2"].alignment = Alignment(horizontal="center")

    # Headers
    headers = ["Unit Type", "Super BUA (sq.ft.)", "Bathrooms", "Units Avail.",
               "Base Rate (Rs./sq.ft.)", "Base Price", "PLC (5%)", "Club + Legal",
               "Approx. Total"]
    header_fill = PatternFill("solid", fgColor="1a1a2e")
    header_font = Font(bold=True, color="FFFFFF", size=10)
    for col, h in enumerate(headers, start=1):
        cell = ws.cell(row=4, column=col, value=h)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
    ws.row_dimensions[4].height = 30

    # Data
    row_fills = [PatternFill("solid", fgColor="f0f0f0"), None]
    for ri, u in enumerate(p["units"]):
        r = ri + 5
        base = u["area"] * u["price_sqft"]
        plc = int(base * 0.05)
        extras = 300000
        total = base + plc + extras

        row_data = [u["type"], u["area"], u["bath"], u["count"],
                    u["price_sqft"],
                    f"Rs. {base/100000:.2f} L",
                    f"Rs. {plc/100000:.2f} L",
                    f"Rs. {extras/100000:.2f} L",
                    lakhs(total)]
        fill = row_fills[ri % 2]
        for ci, val in enumerate(row_data, start=1):
            cell = ws.cell(row=r, column=ci, value=val)
            cell.alignment = Alignment(horizontal="center")
            if fill:
                cell.fill = fill

    # Notes
    note_row = len(p["units"]) + 6
    notes = [
        "* Prices are indicative and subject to change without notice.",
        "* PLC (Preferential Location Charges): Floor rise, park-facing, corner unit extra.",
        "* Extras include: Club membership Rs. 1,50,000 + Legal & Documentation Rs. 1,50,000.",
        "* GST @ 5% applicable on total cost.",
        "* Stamp Duty and Registration extra as per state norms.",
        "* Payment plan details available in the Payment Plan document.",
    ]
    for i, note in enumerate(notes):
        ws.cell(row=note_row + i, column=1, value=note).font = Font(italic=True, size=9, color="666666")

    # Column widths
    from openpyxl.utils import get_column_letter
    col_widths = [22, 18, 12, 14, 20, 18, 16, 16, 20]
    for i, w in enumerate(col_widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w

    path = os.path.join(out_dir, "02_price_list.xlsx")
    wb.save(path)
    print(f"  ✓ {path}")


# =============================================================================
# DOCUMENT 3 — AMENITIES GUIDE (PDF) — ALL PROJECTS
# =============================================================================

def create_amenities_guide(p, out_dir):
    pdf = new_pdf()
    pdf.add_page()

    pdf_h1(pdf, p["name"])
    pdf_h1(pdf, "Amenities & Specifications Guide")
    pdf.set_font("Helvetica", "I", 10)
    pdf.cell(0, 6, p["location"], new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(6)

    pdf_h2(pdf, "Complete Amenities List")
    pdf_bullet(pdf, p["amenities"])

    pdf_h2(pdf, "Construction Specifications")
    specs = [
        "Structure: RCC framed structure with earthquake-resistant design (Zone II)",
        "External walls: 200 mm thick solid concrete block masonry",
        "Internal walls: 100 mm thick solid concrete block masonry",
        "Flooring (living & dining): 800x800 mm vitrified tiles (premium brand)",
        "Flooring (bedrooms): 600x600 mm vitrified tiles",
        "Flooring (kitchen): Anti-skid ceramic tiles",
        "Flooring (bathrooms): Anti-skid ceramic tiles (300x300 mm)",
        "Kitchen: Granite counter top, stainless steel sink, provision for chimney & dishwasher",
        "Bathroom: Branded CP fittings (Jaquar/equivalent), wall-hung WC, shower enclosure in master bath",
        "Doors: Main door — 8-foot teak frame with solid-core flush shutter; internal — hollow-core flush",
        "Windows: UPVC double-glazed sliding windows with mosquito mesh",
        "Electrical: 3-phase supply, concealed copper wiring (Finolex/equivalent), modular switches (Legrand/equivalent)",
        "Lifts: High-speed automatic passenger lifts (Otis/Schindler), 1 service lift per tower",
        "Power backup: 100% DG backup for lifts, common areas, pumps; 750W per apartment",
        "Water: Underground sump + overhead tank, borewell + BWSSB supply",
        "Painting: Interior — JK/Asian Paints premium emulsion; exterior — weatherproof texture paint",
    ]
    pdf_bullet(pdf, specs)

    pdf_h2(pdf, "Green & Sustainability Features")
    green = [
        "Rainwater harvesting system with percolation pits",
        "Sewage treatment plant (STP) — treated water used for landscaping and flushing",
        "LED lighting in all common areas",
        "Solar-powered street lighting",
        "Organic waste converter for wet waste",
        "E-waste collection bins on every floor",
        "EV charging infrastructure pre-installed in parking",
        "Low-VOC paints and adhesives used throughout",
    ]
    pdf_bullet(pdf, green)

    path = os.path.join(out_dir, "03_amenities_guide.pdf")
    pdf.output(path)
    print(f"  ✓ {path}")


# =============================================================================
# DOCUMENT 4 — FAQ DOCUMENT (DOCX) — ALL PROJECTS
# =============================================================================

def create_faq(p, out_dir):
    doc = Document()

    h = doc.add_heading(f"{p['name']} — Frequently Asked Questions", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Location: {p['location']}  |  Developer: {p['developer']}")

    faqs = [
        ("What is the RERA registration number?",
         f"The RERA registration number for {p['short']} is {p['rera']}. "
         f"You can verify this on the Karnataka RERA portal (rera.karnataka.gov.in)."),

        ("When will possession be handed over?",
         f"Possession of {p['short']} is expected by {p['possession']}. "
         f"The construction commenced in {p['construction_start']}. "
         f"RERA-mandated completion date is committed in the registration."),

        ("What unit configurations are available?",
         "The following configurations are available: " +
         ", ".join(f"{u['type']} ({u['area']:,} sq.ft.)" for u in p["units"]) + ". "
         "Each unit is Vastu-compliant and available in multiple floor options."),

        ("What is the price range?",
         f"Base prices range from Rs. {min(u['price_sqft'] for u in p['units']):,}/sq.ft. "
         f"to Rs. {max(u['price_sqft'] for u in p['units']):,}/sq.ft. "
         f"depending on the unit type and floor. The approximate total cost (all-inclusive) ranges from "
         f"{lakhs(min(u['area']*u['price_sqft'] for u in p['units']))} to "
         f"{lakhs(max(u['area']*u['price_sqft'] for u in p['units']))} before GST and registration."),

        ("What amenities are included?",
         f"{p['short']} offers {len(p['amenities'])} premium amenities including: " +
         "; ".join(p["amenities"][:8]) + "; and many more. "
         "Refer to the Amenities Guide document for the complete list."),

        ("Is there a clubhouse and what does it include?",
         f"Yes, {p['short']} has a well-equipped clubhouse with multiple indoor and outdoor facilities "
         f"for residents. Access is through a one-time club membership fee (included in the unit cost)."),

        ("What is the parking policy?",
         "Each residential unit is allotted dedicated covered parking in the basement/podium. "
         "2BHK units get 1 parking slot, 3BHK and above get 2 parking slots. "
         "Visitor parking is provided separately. EV charging points available."),

        ("What is the payment plan?",
         "We offer Construction-Linked Plans (CLP) and Time-Linked Plans (TLP). "
         "A typical CLP structure: 10% booking, 80% linked to construction milestones, 10% on possession. "
         "Home loan tie-ups available with SBI, HDFC, ICICI, Axis Bank, and Kotak Mahindra Bank."),

        ("What is included in the maintenance charges?",
         "Monthly maintenance covers: housekeeping of common areas, landscaping, security, water supply, "
         "power backup for common areas, lift maintenance, and STP operations. "
         "Estimated maintenance: Rs. 3-5 per sq.ft./month (billed quarterly in advance)."),

        ("What approvals does the project have?",
         f"All statutory approvals are in place: RERA registration ({p['rera']}), "
         "BBMP / BDA building plan sanction, environmental clearance, fire NOC, and airport authority clearance. "
         "Title documents are clear and can be reviewed at the sales office."),

        ("What is the floor rise charge?",
         "Floor rise charges are applicable from the 5th floor onwards at Rs. 25-50 per sq.ft. per floor. "
         "Ground floor units attract a ground floor surcharge of Rs. 1,50,000."),

        ("Can I get a home loan for this project?",
         f"{p['short']} is pre-approved by leading banks including SBI, HDFC, ICICI, Axis, and Kotak. "
         "Loan eligibility depends on your income, CIBIL score, and the unit value. "
         "Our in-house team will assist you with loan processing at no extra charge."),

        ("What is the cancellation policy?",
         "If a buyer cancels within 30 days of booking, 2% of the booking amount is deducted. "
         "After 30 days: 10% of the total unit value is forfeited. "
         "Builder cancellation (project delay > 12 months beyond RERA date): full refund + interest at SBI PLR."),

        ("Are there any additional charges beyond the base price?",
         "Yes. Additional charges include: PLC (5% for park-facing / corner / higher floor), "
         "Club membership (Rs. 1,50,000), legal & documentation (Rs. 1,50,000), "
         "GST (5% on total), stamp duty (5% on guideline value), registration (1%), "
         "khata + CC charges (approximately Rs. 50,000)."),

        ("What is the project's connectivity to public transport?",
         f"{p['short']} is well-connected by road and public transport. " +
         " | ".join(p["nearby"][:4]) + ". "
         "BMTC bus stops are within walking distance. Metro connectivity is being expanded in the area."),
    ]

    for i, (q, a) in enumerate(faqs, start=1):
        doc.add_heading(f"Q{i}. {q}", level=2)
        doc.add_paragraph(a)

    path = os.path.join(out_dir, "04_faq.docx")
    doc.save(path)
    print(f"  ✓ {path}")


# =============================================================================
# DOCUMENT 5 — FLOOR PLANS GUIDE (PDF) — FULL SET ONLY
# =============================================================================

def create_floor_plans(p, out_dir):
    pdf = new_pdf()
    pdf.add_page()

    pdf_h1(pdf, p["name"])
    pdf_h1(pdf, "Floor Plans & Unit Layouts Guide")
    pdf.ln(4)

    pdf_h2(pdf, "General Layout Principles")
    pdf_body(pdf,
        f"{p['short']} follows a modern open-plan layout philosophy with maximum natural light and ventilation. "
        f"All apartments are designed to be Vastu-compliant. "
        f"The project has {p['towers']} towers arranged to maximize cross-ventilation and views. "
        f"Floor plates are designed to have a maximum of 4 apartments per floor per tower, "
        "ensuring privacy and exclusivity.")

    for u in p["units"]:
        pdf_h2(pdf, f"{u['type']}  —  {u['area']:,} sq.ft. Super Built-up Area")
        carpet = int(u["area"] * 0.70)
        usable = int(u["area"] * 0.78)
        pdf_kv(pdf, "Super Built-up Area", f"{u['area']:,} sq.ft.")
        pdf_kv(pdf, "Usable Area (approx.)", f"{usable:,} sq.ft.")
        pdf_kv(pdf, "Carpet Area (RERA)", f"{carpet:,} sq.ft.")
        pdf_kv(pdf, "Bathrooms", str(u["bath"]))
        pdf_kv(pdf, "Balconies", "2" if u["area"] >= 1500 else "1")
        pdf_kv(pdf, "Available on Floors",
               f"2nd to {p['floors']}th (preferred: 8th–{p['floors']-2}th for best views)")
        pdf.ln(2)

        # Room dimensions
        pdf.set_font("Helvetica", "B", 10)
        pdf.cell(0, 7, "  Room Dimensions (approximate):", new_x="LMARGIN", new_y="NEXT")
        pdf.set_font("Helvetica", "", 10)

        rooms = []
        if "1BHK" in u["type"]:
            rooms = [("Living + Dining", "14 ft x 12 ft"), ("Bedroom 1", "11 ft x 10 ft"),
                     ("Kitchen", "9 ft x 8 ft"), ("Bathroom", "6 ft x 4.5 ft"), ("Balcony", "9 ft x 4 ft")]
        elif "2BHK" in u["type"]:
            rooms = [("Living Room", "15 ft x 13 ft"), ("Dining Area", "10 ft x 9 ft"),
                     ("Master Bedroom", "14 ft x 12 ft"), ("Bedroom 2", "12 ft x 10 ft"),
                     ("Kitchen", "10 ft x 9 ft"), ("Master Bath", "8 ft x 5 ft"),
                     ("Common Bath", "6 ft x 5 ft"), ("Balcony 1", "10 ft x 4 ft")]
        elif "3BHK" in u["type"]:
            rooms = [("Living Room", "18 ft x 14 ft"), ("Dining Area", "12 ft x 10 ft"),
                     ("Master Bedroom", "16 ft x 13 ft"), ("Bedroom 2", "13 ft x 11 ft"),
                     ("Bedroom 3", "12 ft x 10 ft"), ("Kitchen", "11 ft x 9 ft"),
                     ("Master Bath", "9 ft x 6 ft"), ("Common Bath 1", "7 ft x 5 ft"),
                     ("Common Bath 2", "6 ft x 5 ft"), ("Balcony 1", "12 ft x 4 ft"), ("Balcony 2", "9 ft x 4 ft")]
        else:
            rooms = [("Living Room", "22 ft x 18 ft"), ("Dining Area", "15 ft x 12 ft"),
                     ("Master Bedroom", "18 ft x 15 ft"), ("Bedroom 2", "15 ft x 13 ft"),
                     ("Bedroom 3", "14 ft x 12 ft"), ("Bedroom 4", "13 ft x 11 ft"),
                     ("Kitchen", "14 ft x 11 ft"), ("Study / Puja Room", "9 ft x 8 ft"),
                     ("Master Bath", "11 ft x 7 ft"), ("Balcony 1", "14 ft x 5 ft"), ("Balcony 2", "10 ft x 4 ft")]

        for room, size in rooms:
            pdf.cell(0, 6, f"     {room:<28}  {size}", new_x="LMARGIN", new_y="NEXT")
        pdf.ln(4)

    path = os.path.join(out_dir, "05_floor_plans.pdf")
    pdf.output(path)
    print(f"  ✓ {path}")


# =============================================================================
# DOCUMENT 6 — PAYMENT PLAN (PDF) — FULL SET ONLY
# =============================================================================

def create_payment_plan(p, out_dir):
    pdf = new_pdf()
    pdf.add_page()

    pdf_h1(pdf, p["name"])
    pdf_h1(pdf, "Payment Plans")
    pdf.set_font("Helvetica", "I", 10)
    pdf.cell(0, 6, "Construction-Linked Plan  |  Time-Linked Plan  |  Subvention Plan", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(6)

    # CLP
    pdf_h2(pdf, "Plan A — Construction-Linked Plan (CLP)  [Recommended]")
    pdf_body(pdf, "Pay as the project is built. Lower risk as funds are disbursed based on actual construction milestones verified by bank-appointed valuers.")
    clp_stages = [
        ("On Booking",                                            "10%"),
        ("On Allotment Letter",                                   "10%"),
        ("On Commencement of Foundation",                         "10%"),
        ("On Completion of Ground Floor Slab",                    "10%"),
        ("On Completion of 5th Floor Slab",                      "10%"),
        ("On Completion of 10th Floor Slab",                     "10%"),
        ("On Completion of 15th Floor Slab",                     "5%"),
        ("On Completion of Top Floor Slab",                      "5%"),
        ("On Completion of Brick Work",                          "5%"),
        ("On Completion of Plastering & Internal Finishes",      "5%"),
        ("On Offer of Possession (OC received)",                 "15%"),
        ("On Registration",                                       "5%"),
    ]
    pdf.set_font("Helvetica", "", 10)
    for stage, pct in clp_stages:
        pdf.cell(130, 7, f"  {stage}", new_x="RIGHT", new_y="LAST")
        pdf.cell(0,   7, pct, new_x="LMARGIN", new_y="NEXT", align="R")

    pdf.ln(4)

    # TLP
    pdf_h2(pdf, "Plan B — Time-Linked Plan (TLP)  [2% discount on base price]")
    pdf_body(pdf, "Fixed payment schedule regardless of construction stage. A 2% discount on the base price is offered in exchange for committed payment dates.")
    tlp_stages = [
        ("On Booking",           "20%"),
        ("Within 30 days of Booking", "20%"),
        ("After 6 months",      "20%"),
        ("After 12 months",     "20%"),
        ("After 18 months",     "10%"),
        ("On Possession",       "10%"),
    ]
    for stage, pct in tlp_stages:
        pdf.set_font("Helvetica", "", 10)
        pdf.cell(130, 7, f"  {stage}", new_x="RIGHT", new_y="LAST")
        pdf.cell(0,   7, pct, new_x="LMARGIN", new_y="NEXT", align="R")

    pdf.ln(4)

    # Subvention
    pdf_h2(pdf, "Plan C — Subvention / No EMI Till Possession")
    pdf_body(pdf,
        "Under this scheme, the builder pays the pre-EMI interest on your home loan until possession. "
        "You pay: 10% at booking + 10% within 30 days. The remaining 80% is disbursed via home loan, "
        "and builder covers EMI until OC is received. "
        "Available through our banking partners: SBI, HDFC, ICICI, and Axis Bank only. "
        "Subject to loan sanction and scheme availability.")

    # Home Loans
    pdf_h2(pdf, "Approved Bank Loan Partners")
    banks = [
        ("State Bank of India (SBI)",  "Up to 80% of unit value", "8.50%-9.25% p.a. (floating)"),
        ("HDFC Ltd.",                  "Up to 80% of unit value", "8.70%-9.40% p.a. (floating)"),
        ("ICICI Bank",                 "Up to 75% of unit value", "8.75%-9.50% p.a. (floating)"),
        ("Axis Bank",                  "Up to 75% of unit value", "8.75%-9.30% p.a. (floating)"),
        ("Kotak Mahindra Bank",        "Up to 80% of unit value", "8.65%-9.45% p.a. (floating)"),
    ]
    pdf.set_font("Helvetica", "B", 10)
    pdf.cell(65, 8, "Bank", new_x="RIGHT", new_y="LAST")
    pdf.cell(65, 8, "Max Loan", new_x="RIGHT", new_y="LAST")
    pdf.cell(0,  8, "Rate of Interest", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 10)
    for bank, loan, rate in banks:
        pdf.cell(65, 6, bank, new_x="RIGHT", new_y="LAST")
        pdf.cell(65, 6, loan, new_x="RIGHT", new_y="LAST")
        pdf.cell(0,  6, rate, new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf_body(pdf, "Note: Interest rates are indicative and subject to RBI policy and bank's discretion. "
                  "Final sanction subject to bank's terms and borrower eligibility.")

    path = os.path.join(out_dir, "06_payment_plan.pdf")
    pdf.output(path)
    print(f"  ✓ {path}")


# =============================================================================
# DOCUMENT 7 — LEGAL & RERA SUMMARY (PDF) — FULL SET ONLY
# =============================================================================

def create_legal_rera(p, out_dir):
    pdf = new_pdf()
    pdf.add_page()

    pdf_h1(pdf, p["name"])
    pdf_h1(pdf, "Legal & RERA Compliance Summary")
    pdf.ln(4)

    pdf_h2(pdf, "RERA Registration Details")
    pdf_kv(pdf, "RERA Reg. No.", p["rera"])
    pdf_kv(pdf, "Registered with", "Karnataka Real Estate Regulatory Authority (K-RERA)")
    pdf_kv(pdf, "Portal", "rera.karnataka.gov.in")
    pdf_kv(pdf, "Project Name (RERA)", p["name"])
    pdf_kv(pdf, "Promoter", p["developer"])
    pdf_kv(pdf, "RERA Registration Date", "May 2023")
    pdf_kv(pdf, "RERA Expiry / Completion Date", p["possession"])
    pdf.ln(4)

    pdf_h2(pdf, "Statutory Approvals Obtained")
    approvals = [
        "Building Plan Sanction — BBMP/BDA (obtained)",
        "Environmental Clearance — SEIAA Karnataka (obtained)",
        "Fire NOC — Karnataka Fire & Emergency Services (obtained)",
        "Airport Authority NOC — AAI (obtained, height clearance)",
        "Water & Sewage Connection — BWSSB (obtained)",
        "Electricity Sanction — BESCOM (obtained)",
        "Title Clearance Certificate — verified by empanelled advocate",
        "Encumbrance Certificate — EC obtained and clear",
        "Conversion Order (DC) — agricultural to residential (obtained)",
        "BDA Approved Layout / CDP Zoning — residential zone confirmed",
    ]
    pdf_bullet(pdf, approvals)

    pdf_h2(pdf, "Title & Land Details")
    pdf_body(pdf,
        f"The land on which {p['short']} is being developed is freehold land held by {p['developer']}. "
        "The title is clear and marketable, verified by an independent law firm engaged by the developer. "
        "The land is free from all encumbrances, liens, and litigation. "
        "Survey numbers and khasra numbers are available in the sale agreement and parent documents. "
        "Buyers are advised to engage their own advocate for independent verification before registration.")

    pdf_h2(pdf, "Buyer Rights Under RERA")
    rights = [
        "Right to receive possession on or before the RERA-committed date",
        "Right to receive interest at SBI PLR + 2% on delayed possession (per day)",
        "Right to withdraw from the project and claim full refund with interest in case of material defect",
        "Right to access project documents on RERA portal at any time",
        "Right to form an apartment owners' association (AOA) within 3 months of possession",
        "Right to 5-year structural defect warranty from the developer post-possession",
        "Right to escalation-free pricing — no additional charges beyond those declared in RERA",
    ]
    pdf_bullet(pdf, rights)

    pdf_h2(pdf, "Important Documents Available at Sales Office")
    docs = [
        "Copy of RERA registration certificate",
        "Approved building plan drawings (all floors)",
        "EC (Encumbrance Certificate) for the land",
        "Title deed and parent documents",
        "Environmental clearance order",
        "BBMP/BDA building plan sanction order",
        "Sample sale agreement (ATS and Sale Deed)",
        "List of specifications (schedule of finishes)",
    ]
    pdf_bullet(pdf, docs)

    path = os.path.join(out_dir, "07_legal_rera.pdf")
    pdf.output(path)
    print(f"  ✓ {path}")


# =============================================================================
# DOCUMENT 8 — POSSESSION TIMELINE (XLSX) — FULL SET ONLY
# =============================================================================

def create_possession_timeline(p, out_dir):
    wb = Workbook()
    ws = wb.active
    ws.title = "Possession Timeline"

    ws.merge_cells("A1:F1")
    ws["A1"] = f"{p['name']} — Construction & Possession Timeline"
    ws["A1"].font = Font(bold=True, size=13)
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 24

    headers = ["Milestone", "Description", "Target Month", "Tower(s)", "Status", "Remarks"]
    hfill = PatternFill("solid", fgColor="16213e")
    hfont = Font(bold=True, color="FFFFFF", size=10)
    for ci, h in enumerate(headers, 1):
        cell = ws.cell(row=3, column=ci, value=h)
        cell.fill = hfill
        cell.font = hfont
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    milestones = [
        ("Land & Approvals",    "RERA registration, plan sanction, title clearance",           "Jan–May 2023",    "All",    "Completed", "All approvals in hand"),
        ("Excavation",          "Foundation pit excavation and soil testing",                   "Jun–Aug 2023",    "T1–T4",  "Completed", "Geotechnical report submitted"),
        ("Foundation",          "Pile foundation, raft slab casting",                          "Sep–Dec 2023",    "T1–T4",  "Completed", "100% pile load test done"),
        ("Basement Slab",       "Basement B2, B1 slabs + waterproofing",                       "Jan–Mar 2024",    "T1–T4",  "Completed", "Parking slab ready"),
        ("Ground Floor Slab",   "Podium slab, ground floor framing",                           "Apr–Jun 2024",    "T1–T4",  "In Progress",""),
        ("Floors 1–5",          "Column, beam, slab casting for floors 1–5",                   "Jul–Oct 2024",    "T1–T2",  "In Progress",""),
        ("Floors 6–10",         "Column, beam, slab casting for floors 6–10",                  "Nov 2024–Feb 2025","T1–T2", "Upcoming",  ""),
        ("Floors 11–15",        "Column, beam, slab casting for floors 11–15",                 "Mar–Jun 2025",    "All",    "Upcoming",  ""),
        ("Floors 16–22 (Top)",  "Top floor slab and terrace waterproofing",                    "Jul–Oct 2025",    "All",    "Upcoming",  ""),
        ("Brick Work",          "Internal and external masonry across all towers",              "Aug–Nov 2025",    "All",    "Upcoming",  ""),
        ("Plastering & MEP",    "Internal plaster, electrical, plumbing rough-in",             "Sep 2025–Feb 2026","All",   "Upcoming",  ""),
        ("Internal Finishes",   "Tiling, painting, kitchen, bathroom fixtures",                 "Jan–Jun 2026",    "All",    "Upcoming",  ""),
        ("External Facade",     "External plastering, texture paint, glazing",                 "Mar–Jul 2026",    "All",    "Upcoming",  ""),
        ("Amenities & Landscaping","Clubhouse, pool, gym, gardens fit-out",                   "Jun–Sep 2026",    "Common", "Upcoming",  ""),
        ("Pre-OC Inspection",   "BBMP inspection, fire dept. clearance, lift inspection",      "Oct 2026",        "All",    "Upcoming",  ""),
        ("Occupancy Certificate","OC obtained from BBMP",                                      "Nov 2026",        "All",    "Upcoming",  ""),
        ("Possession Handover",  "Keys handed over to buyers; society formation",              "Dec 2026",        "All",    "Upcoming",  "RERA committed date"),
    ]

    fills = [PatternFill("solid", fgColor="e8f5e9"), PatternFill("solid", fgColor="fff3e0"), None]
    for ri, m in enumerate(milestones):
        r = ri + 4
        for ci, val in enumerate(m, 1):
            cell = ws.cell(row=r, column=ci, value=val)
            cell.alignment = Alignment(wrap_text=True, vertical="top")
            status = m[4]
            if status == "Completed":
                cell.fill = PatternFill("solid", fgColor="c8e6c9")
            elif status == "In Progress":
                cell.fill = PatternFill("solid", fgColor="fff9c4")

    from openpyxl.utils import get_column_letter
    col_widths = [22, 40, 20, 14, 14, 28]
    for i, w in enumerate(col_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w

    for r in range(4, len(milestones) + 4):
        ws.row_dimensions[r].height = 30

    path = os.path.join(out_dir, "08_possession_timeline.xlsx")
    wb.save(path)
    print(f"  ✓ {path}")


# =============================================================================
# MAIN — Generate All Projects
# =============================================================================

def _sanitize_project(p):
    """Recursively replace non-latin-1 chars in all string values of a project dict."""
    import copy
    p = copy.deepcopy(p)

    def clean(obj):
        if isinstance(obj, str):
            return _safe(obj)
        if isinstance(obj, list):
            return [clean(i) for i in obj]
        if isinstance(obj, dict):
            return {k: clean(v) for k, v in obj.items()}
        return obj

    return clean(p)


def main():
    for p in [_sanitize_project(p) for p in PROJECTS]:
        out_dir = os.path.join(BASE_DIR, p["folder"])
        os.makedirs(out_dir, exist_ok=True)
        print(f"\n📁 {p['folder']} — {p['name']}")

        create_brochure(p, out_dir)
        create_price_list(p, out_dir)
        create_amenities_guide(p, out_dir)
        create_faq(p, out_dir)

        if p.get("full_set"):
            create_floor_plans(p, out_dir)
            create_payment_plan(p, out_dir)
            create_legal_rera(p, out_dir)
            create_possession_timeline(p, out_dir)

    print("\n" + "=" * 60)
    print("✅  All project documents generated successfully!")
    print(f"📂  Output folder: {os.path.abspath(BASE_DIR)}")
    print("=" * 60)


if __name__ == "__main__":
    main()
