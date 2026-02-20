"""Generate real estate documents for RAG testing."""

import os
from fpdf import FPDF
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT_DIR = "./real_estate_documents"
os.makedirs(OUTPUT_DIR, exist_ok=True)


# ──────────────────────────────────────────────
# PDF 1: Sunrise Heights - Premium Apartments (3-4 pages)
# ──────────────────────────────────────────────
def create_sunrise_heights():
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Page 1 - Overview
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 22)
    pdf.cell(0, 15, "Sunrise Heights Premium Apartments", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 12)
    pdf.cell(0, 8, "Whitefield, Bangalore - 560066", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(10)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "Project Overview", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    overview = (
        "Sunrise Heights is a premium residential project spread across 5.5 acres of lush green landscape "
        "in the heart of Whitefield, Bangalore. Developed by Prestige Group, this project offers 2BHK, 3BHK, "
        "and 4BHK apartments designed with modern architecture and world-class amenities. The project comprises "
        "4 towers with 22 floors each, housing a total of 520 residential units. Construction started in January "
        "2023 and possession is expected by December 2026.\n\n"
        "The project is strategically located near the Whitefield Metro Station (1.2 km), ITPL Tech Park (2.5 km), "
        "and Phoenix Marketcity Mall (3 km). Major hospitals including Manipal Hospital (1.8 km) and Columbia Asia "
        "(2.2 km) are within close proximity. International schools such as Inventure Academy (1.5 km) and "
        "Whitefield Global School (2 km) are nearby."
    )
    pdf.multi_cell(0, 6, overview)
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 14)
    pdf.cell(0, 10, "Pricing Details", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)

    pricing = [
        ["Unit Type", "Super Built-Up Area", "Base Price", "Total Price (Approx.)"],
        ["2 BHK Standard", "1,150 sq.ft.", "Rs. 7,500/sq.ft.", "Rs. 86.25 Lakhs"],
        ["2 BHK Premium", "1,280 sq.ft.", "Rs. 7,800/sq.ft.", "Rs. 99.84 Lakhs"],
        ["3 BHK Standard", "1,650 sq.ft.", "Rs. 7,500/sq.ft.", "Rs. 1.24 Crores"],
        ["3 BHK Premium", "1,850 sq.ft.", "Rs. 7,800/sq.ft.", "Rs. 1.44 Crores"],
        ["3 BHK Duplex", "2,200 sq.ft.", "Rs. 8,200/sq.ft.", "Rs. 1.80 Crores"],
        ["4 BHK Penthouse", "3,500 sq.ft.", "Rs. 9,000/sq.ft.", "Rs. 3.15 Crores"],
    ]

    col_widths = [40, 45, 40, 50]
    pdf.set_font("Helvetica", "B", 10)
    for j, cell in enumerate(pricing[0]):
        pdf.cell(col_widths[j], 8, cell, border=1, align="C")
    pdf.ln()
    pdf.set_font("Helvetica", "", 10)
    for row in pricing[1:]:
        for j, cell in enumerate(row):
            pdf.cell(col_widths[j], 7, cell, border=1, align="C")
        pdf.ln()

    pdf.ln(5)
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Additional charges: Car parking - Rs. 3.5 Lakhs per slot (one complimentary with 3BHK and above). "
        "Club house membership - Rs. 2 Lakhs. Legal and registration charges as per government norms. "
        "GST of 5% applicable on under-construction properties. Maintenance deposit of Rs. 50 per sq.ft. "
        "collected at the time of possession."
    )

    # Page 2 - Amenities
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, "Amenities & Facilities", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Sports & Fitness", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "- Temperature-controlled swimming pool (25m lap pool + kids pool)\n"
        "- Fully equipped gymnasium with Technogym equipment (3,500 sq.ft.)\n"
        "- Tennis court (synthetic surface, floodlit)\n"
        "- Badminton court (2 indoor courts)\n"
        "- Basketball court (half court)\n"
        "- Cricket practice nets with bowling machine\n"
        "- Squash court\n"
        "- Table tennis room (2 tables)\n"
        "- Jogging track (600m rubberized track around the perimeter)\n"
        "- Yoga and aerobics studio (1,200 sq.ft.)\n"
        "- Cycling track dedicated for residents"
    )
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Lifestyle & Recreation", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "- Clubhouse with banquet hall (5,000 sq.ft., capacity 200 guests)\n"
        "- Mini theater / screening room (50 seats, Dolby Atmos sound)\n"
        "- Indoor games room (billiards, foosball, carrom, chess)\n"
        "- Library and reading lounge with 2,000+ books\n"
        "- Co-working space with high-speed internet (20 workstations)\n"
        "- Party lawn and barbecue area\n"
        "- Rooftop sky lounge with panoramic city views\n"
        "- Spa and sauna facility\n"
        "- Meditation garden with water features\n"
        "- Pet park with agility equipment"
    )
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Children's Facilities", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "- Children's play area with imported play equipment (ages 2-5 and 6-12)\n"
        "- Splash pad and water play zone\n"
        "- Indoor play zone with soft play area\n"
        "- Art and craft room\n"
        "- Day care center (professionally managed)\n"
        "- Outdoor adventure zone with climbing wall"
    )

    # Page 3 - Security & Specifications
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, "Security & Technical Specifications", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Security Features", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "- 3-tier security system with boom barriers at entry/exit\n"
        "- 24/7 CCTV surveillance with 200+ cameras (AI-powered analytics)\n"
        "- Video door phone in every apartment connected to guard room\n"
        "- Biometric access for lobbies and common areas\n"
        "- Intercom facility connecting all apartments to security\n"
        "- Fire detection and suppression system (sprinklers in all common areas)\n"
        "- Earthquake-resistant RCC framed structure (Zone II compliant)\n"
        "- Lightning arrestor on all towers\n"
        "- 100% DG power backup for common areas, 5 kVA per apartment"
    )
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Construction Specifications", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Flooring: Italian marble in living and dining areas; vitrified tiles in bedrooms (800x800mm); "
        "anti-skid ceramic tiles in bathrooms and balconies; hardwood laminate in master bedroom (optional upgrade).\n\n"
        "Kitchen: Modular kitchen with granite countertop, stainless steel sink, ceramic tile dado up to 2 feet "
        "above countertop. Provision for water purifier, dishwasher, and chimney. Piped gas connection from "
        "central gas bank.\n\n"
        "Bathrooms: Premium sanitary ware from Kohler/Grohe. Hot and cold water mixer. Concealed plumbing. "
        "Anti-skid tiles. Frameless glass shower enclosure in master bathroom.\n\n"
        "Doors: Main door - teak wood frame with veneer shutter and digital smart lock. Internal doors - "
        "flush doors with veneer finish. French windows in living room with UPVC frames and double-glazed glass.\n\n"
        "Electrical: Modular switches from Schneider/Legrand. Concealed copper wiring. AC provision in all "
        "bedrooms and living room. EV charging provision in parking. Smart home automation ready."
    )

    # Page 4 - Floor Plans & Payment
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 12, "Payment Plans & Contact", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(5)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Flexi Payment Plan", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Milestone 1: 10% on booking (within 30 days of application)\n"
        "Milestone 2: 15% on execution of agreement\n"
        "Milestone 3: 10% on completion of foundation\n"
        "Milestone 4: 10% on completion of 5th floor slab\n"
        "Milestone 5: 10% on completion of 12th floor slab\n"
        "Milestone 6: 10% on completion of 20th floor slab\n"
        "Milestone 7: 10% on completion of brickwork\n"
        "Milestone 8: 10% on completion of flooring\n"
        "Milestone 9: 10% on completion of final finishes\n"
        "Milestone 10: 5% on possession / handover"
    )
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Bank Loan Approvals", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "The project is pre-approved by all major banks including SBI, HDFC, ICICI, Axis Bank, Bank of Baroda, "
        "Punjab National Bank, Kotak Mahindra Bank, and LIC Housing Finance. Home loan interest rates starting "
        "from 8.35% per annum. The project is RERA registered under RERA No. PRM/KA/RERA/1251/310/PR/190524/006456."
    )
    pdf.ln(3)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Sales Office & Contact", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Sales Office: Plot No. 45, ITPL Main Road, Whitefield, Bangalore - 560066\n"
        "Timings: 10:00 AM to 7:00 PM (All days including weekends)\n"
        "Phone: +91 80 4567 8900\n"
        "Email: sales@sunriseheights.in\n"
        "Website: www.sunriseheights.in\n\n"
        "Site visits available by appointment. Complimentary pickup from Whitefield Metro Station."
    )

    pdf.output(os.path.join(OUTPUT_DIR, "01_Sunrise_Heights_Premium_Apartments.pdf"))
    print("Created: 01_Sunrise_Heights_Premium_Apartments.pdf (4 pages)")


# ──────────────────────────────────────────────
# PDF 2: Green Valley Villas
# ──────────────────────────────────────────────
def create_green_valley():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 12, "Green Valley Luxury Villas", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, "Sarjapur Road, Bangalore - 562125", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Project Details", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Green Valley is an exclusive gated community of 85 independent luxury villas spread across 25 acres "
        "with 70% open space. Each villa is designed with Vastu-compliant architecture and offers private gardens, "
        "dedicated parking for 2 cars, and terrace access.\n\n"
        "Villa Types:\n"
        "- 3 BHK Villa: 2,400 sq.ft. built-up on 1,200 sq.ft. plot - Rs. 1.95 Crores\n"
        "- 4 BHK Villa: 3,200 sq.ft. built-up on 1,500 sq.ft. plot - Rs. 2.85 Crores\n"
        "- 4 BHK Duplex Villa: 4,000 sq.ft. built-up on 2,400 sq.ft. plot - Rs. 3.75 Crores\n"
        "- 5 BHK Presidential Villa: 5,500 sq.ft. built-up on 3,600 sq.ft. plot - Rs. 5.50 Crores\n\n"
        "Key Amenities:\n"
        "- Private clubhouse with infinity pool, gym, and party hall\n"
        "- 9-hole mini golf course\n"
        "- Organic vegetable garden (dedicated plots for each villa)\n"
        "- Amphitheater for community events\n"
        "- Senior citizen park with reflexology pathway\n"
        "- Rainwater harvesting and solar panels on all villas\n"
        "- Sewage treatment plant with recycled water for landscaping\n"
        "- 24/7 concierge service\n"
        "- Gated compound with 3-level security\n\n"
        "Location Advantages: Sarjapur Road is one of Bangalore's fastest-growing corridors. "
        "Close to major IT parks (Wipro SEZ 3 km, Infosys Campus 5 km), Decathlon (2 km), "
        "Total Mall (1.5 km), and multiple international schools."
    )

    pdf.output(os.path.join(OUTPUT_DIR, "02_Green_Valley_Luxury_Villas.pdf"))
    print("Created: 02_Green_Valley_Luxury_Villas.pdf")


# ──────────────────────────────────────────────
# PDF 3: Metro Edge Compact Homes
# ──────────────────────────────────────────────
def create_metro_edge():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 12, "Metro Edge Compact Homes", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, "Electronic City Phase 2, Bangalore - 560100", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)

    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Metro Edge offers affordable smart homes designed for young professionals and first-time homebuyers. "
        "Located just 400 meters from the upcoming Electronic City Metro Station, this project ensures excellent "
        "connectivity to the entire city.\n\n"
        "Unit Configuration & Pricing:\n"
        "- 1 BHK Smart Home: 550 sq.ft. - Rs. 32 Lakhs\n"
        "- 1 BHK + Study: 650 sq.ft. - Rs. 38 Lakhs\n"
        "- 2 BHK Compact: 850 sq.ft. - Rs. 49 Lakhs\n"
        "- 2 BHK Standard: 1,050 sq.ft. - Rs. 58 Lakhs\n"
        "- 3 BHK: 1,350 sq.ft. - Rs. 74 Lakhs\n\n"
        "All units come with smart home features including app-controlled lighting, Alexa-compatible switches, "
        "smart door locks, and video doorbell.\n\n"
        "Project Features:\n"
        "- 3 towers, 30 floors each, 720 units total\n"
        "- Rooftop solar panels reducing common area electricity by 40%\n"
        "- EV charging stations on every parking level\n"
        "- Co-working lounge with meeting rooms on the ground floor\n"
        "- Rooftop infinity pool with city views\n"
        "- Sky gym on the 28th floor\n"
        "- Multipurpose court (basketball/volleyball)\n"
        "- Indoor squash court\n"
        "- Children's play area and creche\n"
        "- Convenience store and pharmacy within the complex\n"
        "- ATM facility on premises\n\n"
        "Maintenance charges: Rs. 3.50 per sq.ft. per month. Sinking fund: Rs. 30 per sq.ft. one-time."
    )

    pdf.output(os.path.join(OUTPUT_DIR, "03_Metro_Edge_Compact_Homes.pdf"))
    print("Created: 03_Metro_Edge_Compact_Homes.pdf")


# ──────────────────────────────────────────────
# PDF 4: Royal Orchid Residency
# ──────────────────────────────────────────────
def create_royal_orchid():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 12, "Royal Orchid Residency", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, "Hebbal, Bangalore - 560024", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)

    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Royal Orchid Residency is a luxury high-rise project located near Hebbal Lake, offering stunning "
        "lake views from select apartments. Developed by Brigade Group, this iconic 45-floor tower is one "
        "of the tallest residential buildings in North Bangalore.\n\n"
        "Unit Types & Pricing:\n"
        "- 3 BHK Lake View: 1,800 sq.ft. - Rs. 1.62 Crores (Rs. 9,000/sq.ft.)\n"
        "- 3 BHK City View: 1,800 sq.ft. - Rs. 1.44 Crores (Rs. 8,000/sq.ft.)\n"
        "- 4 BHK Sky Suite: 2,800 sq.ft. - Rs. 2.80 Crores (Rs. 10,000/sq.ft.)\n"
        "- 4 BHK Penthouse: 4,200 sq.ft. - Rs. 5.04 Crores (Rs. 12,000/sq.ft.)\n\n"
        "Premium Amenities:\n"
        "- Helipad on rooftop (for emergency medical evacuation)\n"
        "- Sky lounge and bar on 43rd floor\n"
        "- Observation deck on 44th floor with telescope\n"
        "- Temperature-controlled infinity pool on 15th floor podium\n"
        "- Professional-grade tennis court with coaching facility\n"
        "- Spa with steam, sauna, and jacuzzi\n"
        "- Business center with video conferencing facility\n"
        "- Private dining room for residents (chef on call)\n"
        "- Wine cellar and cigar lounge\n"
        "- Pet grooming station\n"
        "- Concierge desk with travel and event planning services\n\n"
        "Connectivity: 2 km from Hebbal Flyover, 8 km from Kempegowda International Airport, "
        "direct access to Outer Ring Road and Bellary Road. Manyata Tech Park 3 km, "
        "Esteem Mall 1 km, Columbia Asia Hospital 1.5 km."
    )

    pdf.output(os.path.join(OUTPUT_DIR, "04_Royal_Orchid_Residency.pdf"))
    print("Created: 04_Royal_Orchid_Residency.pdf")


# ──────────────────────────────────────────────
# PDF 5: Eco Habitat Township
# ──────────────────────────────────────────────
def create_eco_habitat():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 12, "Eco Habitat Integrated Township", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, "Devanahalli, Bangalore - 562110", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)

    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Eco Habitat is a 100-acre integrated township near Bangalore International Airport. The township "
        "includes residential zones (apartments, villas, row houses), a commercial plaza, international school, "
        "hospital, and a 15-acre central park.\n\n"
        "Residential Options:\n"
        "- Studio Apartment: 450 sq.ft. - Rs. 22 Lakhs\n"
        "- 1 BHK Apartment: 650 sq.ft. - Rs. 32 Lakhs\n"
        "- 2 BHK Apartment: 1,100 sq.ft. - Rs. 52 Lakhs\n"
        "- 3 BHK Apartment: 1,500 sq.ft. - Rs. 72 Lakhs\n"
        "- Row House (3 BHK): 1,800 sq.ft. on 1,000 sq.ft. plot - Rs. 1.15 Crores\n"
        "- Villa (4 BHK): 3,000 sq.ft. on 2,000 sq.ft. plot - Rs. 2.10 Crores\n\n"
        "Eco-Friendly Features:\n"
        "- IGBC Platinum certified (Indian Green Building Council)\n"
        "- 100% solar-powered common areas\n"
        "- Rainwater harvesting with storage capacity of 5 lakh liters\n"
        "- Greywater recycling for landscaping and flushing\n"
        "- Organic waste composting plant\n"
        "- Electric shuttle service within township\n"
        "- Bicycle-sharing program with 100 cycles\n"
        "- Vertical gardens on all apartment buildings\n"
        "- Butterfly garden and bird sanctuary (2 acres)\n\n"
        "Township Facilities:\n"
        "- 10-acre sports complex with Olympic-size pool, athletics track, and indoor stadium\n"
        "- 2-screen multiplex cinema\n"
        "- Retail plaza with supermarket, restaurants, and banks\n"
        "- 200-bed multi-specialty hospital\n"
        "- K-12 international school (CBSE and IB curriculum)\n"
        "- Fire station within the township\n\n"
        "Expected completion: Phase 1 (apartments) - June 2025, Phase 2 (villas) - December 2026, "
        "Phase 3 (commercial) - March 2027."
    )

    pdf.output(os.path.join(OUTPUT_DIR, "05_Eco_Habitat_Township.pdf"))
    print("Created: 05_Eco_Habitat_Township.pdf")


# ──────────────────────────────────────────────
# PDF 6: Lakeshore Towers
# ──────────────────────────────────────────────
def create_lakeshore():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 12, "Lakeshore Towers", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, "Bellandur, Bangalore - 560103", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)

    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Lakeshore Towers is a waterfront residential project overlooking Bellandur Lake. The project "
        "offers a unique lakeside living experience with private lake-facing balconies and dedicated "
        "promenade walkway.\n\n"
        "Configuration:\n"
        "- 2 BHK Lake View: 1,200 sq.ft. - Rs. 96 Lakhs\n"
        "- 3 BHK Lake View: 1,750 sq.ft. - Rs. 1.40 Crores\n"
        "- 3 BHK Corner Unit: 1,900 sq.ft. - Rs. 1.52 Crores\n"
        "- 4 BHK Premium: 2,600 sq.ft. - Rs. 2.34 Crores\n\n"
        "Amenities:\n"
        "- Lakeside promenade (800m walking/cycling track)\n"
        "- Kayaking and pedal boats on the lake\n"
        "- Outdoor amphitheater facing the lake\n"
        "- Sunrise yoga deck by the waterfront\n"
        "- Infinity pool with lake views\n"
        "- Open-air gymnasium\n"
        "- Bird watching station\n"
        "- Fishing pier (catch and release)\n"
        "- Lakeside cafe and restaurant\n"
        "- Bonfire and barbecue area\n\n"
        "Specifications: Earthquake-resistant structure, VRV air conditioning, imported marble flooring, "
        "German kitchen fittings (Hacker), Italian bathroom fixtures (Villeroy & Boch), "
        "smart home automation (Schneider Wiser), double-glazed windows for sound insulation.\n\n"
        "Monthly maintenance: Rs. 4.50 per sq.ft. Possession: December 2025. "
        "RERA No: PRM/KA/RERA/1251/309/PR/200623/005892"
    )

    pdf.output(os.path.join(OUTPUT_DIR, "06_Lakeshore_Towers.pdf"))
    print("Created: 06_Lakeshore_Towers.pdf")


# ──────────────────────────────────────────────
# PDF 7: Heritage Grand - Senior Living
# ──────────────────────────────────────────────
def create_heritage_grand():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 12, "Heritage Grand - Senior Living", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, "Yelahanka, Bangalore - 560064", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)

    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Heritage Grand is Bangalore's premier senior living community designed exclusively for residents "
        "aged 55 and above. Spread across 8 acres of serene, landscaped grounds with 24/7 medical support.\n\n"
        "Living Options:\n"
        "- 1 BHK Independent Unit: 700 sq.ft. - Rs. 45 Lakhs\n"
        "- 2 BHK Independent Unit: 1,000 sq.ft. - Rs. 65 Lakhs\n"
        "- 2 BHK Premium Unit: 1,200 sq.ft. - Rs. 78 Lakhs\n"
        "- 3 BHK Family Unit: 1,500 sq.ft. - Rs. 95 Lakhs\n\n"
        "Healthcare & Wellness:\n"
        "- On-site clinic with resident doctor (24/7)\n"
        "- Tie-up with Manipal Hospital for emergencies (ambulance within 10 minutes)\n"
        "- Physiotherapy center\n"
        "- Memory care wing for dementia patients\n"
        "- Emergency call buttons in every room and bathroom\n"
        "- Anti-skid flooring throughout\n"
        "- Wheelchair-accessible design with ramps and wide corridors\n"
        "- Medicine management and reminder service\n\n"
        "Community Features:\n"
        "- Community dining hall with nutritionist-planned meals (included in maintenance)\n"
        "- Temple, prayer hall, and meditation room\n"
        "- Library with 5,000 books and daily newspapers\n"
        "- Arts and crafts studio\n"
        "- Music room with instruments\n"
        "- Heated swimming pool with gentle entry\n"
        "- Walking paths with shaded seating every 50 meters\n"
        "- Vegetable garden plots for residents\n"
        "- Weekly cultural programs and movie screenings\n"
        "- Housekeeping and laundry service included\n\n"
        "Monthly maintenance (includes meals, housekeeping, security): Rs. 15,000 for 1 BHK, "
        "Rs. 20,000 for 2 BHK, Rs. 25,000 for 3 BHK."
    )

    pdf.output(os.path.join(OUTPUT_DIR, "07_Heritage_Grand_Senior_Living.pdf"))
    print("Created: 07_Heritage_Grand_Senior_Living.pdf")


# ──────────────────────────────────────────────
# PDF 8: Skyline Commercial Plaza
# ──────────────────────────────────────────────
def create_skyline_commercial():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 12, "Skyline Commercial Plaza", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, "Outer Ring Road, Marathahalli, Bangalore - 560037", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)

    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Skyline Commercial Plaza is a Grade-A commercial development offering office spaces, retail shops, "
        "and food court on Outer Ring Road, Bangalore's prime commercial corridor.\n\n"
        "Available Spaces:\n"
        "- Retail Shops (Ground Floor): 300-1,500 sq.ft. - Rs. 15,000/sq.ft.\n"
        "- Office Space (Floors 2-8): 500-5,000 sq.ft. - Rs. 10,500/sq.ft.\n"
        "- Premium Office (Floors 9-12): 1,000-3,000 sq.ft. - Rs. 12,000/sq.ft.\n"
        "- Food Court Units (Floor 1): 200-800 sq.ft. - Rs. 18,000/sq.ft.\n"
        "- Penthouse Office (Floor 13): 8,000 sq.ft. - Rs. 14,000/sq.ft.\n\n"
        "Commercial Amenities:\n"
        "- 6 high-speed elevators (8 persons each)\n"
        "- Central air conditioning\n"
        "- 500 car basement parking + 200 two-wheeler slots\n"
        "- 100% power backup with DG sets\n"
        "- Fiber optic internet infrastructure\n"
        "- Food court with seating for 300\n"
        "- Conference center with 3 meeting rooms\n"
        "- ATMs from 4 major banks\n"
        "- Cafeteria on terrace level\n"
        "- LEED Gold certification\n\n"
        "Rental yields: Expected 7-9% annual returns. Current market rent: Rs. 55-75 per sq.ft. per month "
        "for office spaces on Outer Ring Road."
    )

    pdf.output(os.path.join(OUTPUT_DIR, "08_Skyline_Commercial_Plaza.pdf"))
    print("Created: 08_Skyline_Commercial_Plaza.pdf")


# ──────────────────────────────────────────────
# PDF 9: Palm Grove Plots
# ──────────────────────────────────────────────
def create_palm_grove():
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 12, "Palm Grove Plotted Development", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, "Kanakapura Road, Bangalore - 560082", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)

    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Palm Grove is a BMRDA-approved plotted development spread across 50 acres on Kanakapura Road, "
        "one of Bangalore's rapidly developing southern corridors.\n\n"
        "Plot Sizes & Pricing:\n"
        "- 1,200 sq.ft. (30x40) - Rs. 48 Lakhs (Rs. 4,000/sq.ft.)\n"
        "- 1,500 sq.ft. - Rs. 60 Lakhs (Rs. 4,000/sq.ft.)\n"
        "- 2,400 sq.ft. (30x80) - Rs. 96 Lakhs (Rs. 4,000/sq.ft.)\n"
        "- 4,000 sq.ft. (40x100) - Rs. 1.72 Crores (Rs. 4,300/sq.ft.)\n"
        "- Corner plots carry 10% premium\n"
        "- Park-facing plots carry 5% premium\n\n"
        "Infrastructure:\n"
        "- 40-foot main roads and 30-foot internal roads (asphalted)\n"
        "- Underground drainage and sewage system\n"
        "- Bore well water supply with overhead tank\n"
        "- BESCOM electricity connection with transformer\n"
        "- Street lighting with solar-powered LED lights\n"
        "- Piped gas provision\n"
        "- Compound wall around the entire layout\n"
        "- Landscaped median on main roads\n\n"
        "Common Amenities:\n"
        "- 5,000 sq.ft. clubhouse with gym and indoor games\n"
        "- Swimming pool\n"
        "- Children's play area\n"
        "- Jogging track (1.2 km)\n"
        "- 3-acre central park with amphitheater\n"
        "- Basketball court\n"
        "- Gated community with 24/7 security\n\n"
        "Location: 4 km from NICE Road junction, 12 km from JP Nagar, 18 km from MG Road. "
        "Upcoming metro line extension will have a station 2 km from the project."
    )

    pdf.output(os.path.join(OUTPUT_DIR, "09_Palm_Grove_Plots.pdf"))
    print("Created: 09_Palm_Grove_Plots.pdf")


# ──────────────────────────────────────────────
# PDF 10: Bangalore Real Estate Market Report
# ──────────────────────────────────────────────
def create_market_report():
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 20)
    pdf.cell(0, 12, "Bangalore Real Estate Market Report 2024-25", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.set_font("Helvetica", "", 11)
    pdf.cell(0, 8, "Annual Analysis by PropertyInsights Research", new_x="LMARGIN", new_y="NEXT", align="C")
    pdf.ln(8)

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Market Overview", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Bangalore's real estate market continued its upward trajectory in 2024-25, with residential property "
        "prices appreciating by 12-18% across major micro-markets. The city saw approximately 62,000 new "
        "residential unit launches, a 15% increase over the previous year. Total sales volume reached 58,000 "
        "units, indicating healthy absorption rates.\n\n"
        "Key Market Indicators:\n"
        "- Average residential price: Rs. 6,800 per sq.ft. (city average)\n"
        "- Highest appreciation corridor: Whitefield (18% YoY)\n"
        "- Most affordable corridor: Devanahalli-Airport Road (Rs. 4,500/sq.ft. average)\n"
        "- Most expensive micro-market: Koramangala (Rs. 14,000/sq.ft. average)\n"
        "- Unsold inventory: 45,000 units (8.5 months supply)\n"
        "- Rental yield average: 3.2% for residential, 7.5% for commercial\n\n"
    )

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Area-wise Price Analysis (Rs. per sq.ft.)", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "Koramangala: Rs. 12,000 - 16,000 | HSR Layout: Rs. 8,500 - 11,000\n"
        "Whitefield: Rs. 6,500 - 9,000 | Electronic City: Rs. 4,800 - 6,500\n"
        "Sarjapur Road: Rs. 5,500 - 8,000 | Hebbal: Rs. 7,500 - 10,000\n"
        "Yelahanka: Rs. 5,000 - 7,000 | Bannerghatta Road: Rs. 5,800 - 8,500\n"
        "Kanakapura Road: Rs. 3,800 - 5,500 | Devanahalli: Rs. 3,500 - 5,000\n"
        "Marathahalli: Rs. 7,000 - 9,500 | Bellandur: Rs. 7,500 - 10,500\n"
        "Indiranagar: Rs. 13,000 - 18,000 | Jayanagar: Rs. 10,000 - 14,000\n\n"
    )

    pdf.set_font("Helvetica", "B", 13)
    pdf.cell(0, 10, "Emerging Trends", new_x="LMARGIN", new_y="NEXT")
    pdf.set_font("Helvetica", "", 11)
    pdf.multi_cell(0, 6,
        "1. Branded Residences: Luxury developers partnering with hospitality brands (Four Seasons, "
        "Ritz-Carlton) for branded residences priced at Rs. 15,000-25,000 per sq.ft.\n\n"
        "2. Co-living Spaces: Growing demand from IT professionals. Major players like Zolo, CoLive "
        "expanding rapidly. Average rent: Rs. 12,000-18,000 per month.\n\n"
        "3. Sustainable Buildings: 35% of new launches in 2024 had green certifications, up from 20% "
        "in 2022. Buyers willing to pay 5-8% premium for eco-friendly features.\n\n"
        "4. Smart Homes: 60% of new premium projects (above Rs. 1 Crore) offering smart home features "
        "as standard. Home automation market growing at 25% CAGR.\n\n"
        "5. Peripheral Growth: North Bangalore (Devanahalli-Yelahanka) and East Bangalore (Budigere-Whitefield) "
        "seeing highest new project launches due to airport connectivity and IT park proximity."
    )

    pdf.output(os.path.join(OUTPUT_DIR, "10_Bangalore_Market_Report_2024.pdf"))
    print("Created: 10_Bangalore_Market_Report_2024.pdf")


# ──────────────────────────────────────────────
# XLSX 1: Property Comparison Sheet
# ──────────────────────────────────────────────
def create_property_comparison_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Property Comparison"

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="0D6CF2", end_color="0D6CF2", fill_type="solid")
    border = Border(
        left=Side(style="thin"), right=Side(style="thin"),
        top=Side(style="thin"), bottom=Side(style="thin"),
    )

    headers = [
        "Project Name", "Location", "Developer", "Type", "1 BHK Price",
        "2 BHK Price", "3 BHK Price", "4 BHK Price", "Total Units",
        "Floors", "Possession Date", "RERA Registered", "Amenities Count",
        "Parking", "Maintenance (Rs/sqft/month)"
    ]

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", wrap_text=True)
        cell.border = border

    data = [
        ["Sunrise Heights", "Whitefield", "Prestige Group", "Apartment", "N/A", "86.25L", "1.24Cr", "3.15Cr", 520, 22, "Dec 2026", "Yes", 35, "1 per 2BHK, 2 per 3BHK+", 4.00],
        ["Green Valley Villas", "Sarjapur Road", "Sobha Ltd", "Villa", "N/A", "N/A", "1.95Cr", "2.85Cr", 85, 2, "Mar 2026", "Yes", 20, "2 per villa", 5.50],
        ["Metro Edge", "Electronic City", "Godrej Properties", "Apartment", "32L", "49L", "74L", "N/A", 720, 30, "Sep 2025", "Yes", 28, "1 per unit", 3.50],
        ["Royal Orchid", "Hebbal", "Brigade Group", "Apartment", "N/A", "N/A", "1.44Cr", "2.80Cr", 320, 45, "Jun 2026", "Yes", 40, "2 per unit", 5.00],
        ["Eco Habitat", "Devanahalli", "Puravankara", "Township", "22L (Studio)", "52L", "72L", "2.10Cr (Villa)", 1200, "Varies", "Phase-wise", "Yes", 50, "1-2 per unit", 3.00],
        ["Lakeshore Towers", "Bellandur", "Prestige Group", "Apartment", "N/A", "96L", "1.40Cr", "2.34Cr", 280, 28, "Dec 2025", "Yes", 25, "1 per 2BHK, 2 per 3BHK+", 4.50],
        ["Heritage Grand", "Yelahanka", "Columbia Pacific", "Senior Living", "45L", "65L", "95L", "N/A", 150, 4, "Ready", "Yes", 30, "1 per unit", "15,000 flat"],
        ["Skyline Plaza", "Marathahalli", "Embassy Group", "Commercial", "N/A", "N/A", "N/A", "N/A", 200, 13, "Mar 2025", "Yes", 15, "500 cars", 6.00],
        ["Palm Grove", "Kanakapura Rd", "Total Environment", "Plots", "N/A", "N/A", "N/A", "N/A", 350, "N/A", "Jun 2025", "Yes", 12, "N/A", 2.00],
        ["Prestige Lakeside", "Varthur", "Prestige Group", "Apartment", "N/A", "72L", "1.05Cr", "1.85Cr", 450, 18, "Mar 2027", "Yes", 32, "1 per unit", 3.75],
        ["Brigade Cornerstone", "Whitefield", "Brigade Group", "Apartment", "N/A", "65L", "92L", "N/A", 600, 15, "Dec 2025", "Yes", 25, "1 per unit", 3.25],
        ["Sobha Dream Acres", "Panathur", "Sobha Ltd", "Apartment", "N/A", "55L", "82L", "1.20Cr", 800, 20, "Ready", "Yes", 30, "1-2 per unit", 3.50],
    ]

    for row_idx, row_data in enumerate(data, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.border = border
            cell.alignment = Alignment(horizontal="center", wrap_text=True)

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[ws.cell(row=1, column=col).column_letter].width = 18

    # Sheet 2: Amenities Matrix
    ws2 = wb.create_sheet("Amenities Matrix")
    amenities = [
        "Swimming Pool", "Gym", "Tennis Court", "Clubhouse", "Jogging Track",
        "Children's Play Area", "Library", "Co-working Space", "Spa/Sauna",
        "Mini Theater", "Basketball Court", "Party Hall", "Garden/Park",
        "EV Charging", "Smart Home", "Solar Power", "Rainwater Harvesting",
        "Pet Park", "Indoor Games", "Security (CCTV)"
    ]
    projects = ["Sunrise Heights", "Green Valley", "Metro Edge", "Royal Orchid",
                "Eco Habitat", "Lakeshore", "Heritage Grand"]

    ws2.cell(row=1, column=1, value="Amenity").font = header_font
    ws2.cell(row=1, column=1).fill = header_fill
    for j, proj in enumerate(projects, 2):
        cell = ws2.cell(row=1, column=j, value=proj)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    import random
    random.seed(42)
    for i, amenity in enumerate(amenities, 2):
        ws2.cell(row=i, column=1, value=amenity).font = Font(bold=True)
        for j in range(2, len(projects) + 2):
            val = "Yes" if random.random() > 0.25 else "No"
            ws2.cell(row=i, column=j, value=val).alignment = Alignment(horizontal="center")

    ws2.column_dimensions["A"].width = 22
    for col in range(2, len(projects) + 2):
        ws2.column_dimensions[ws2.cell(row=1, column=col).column_letter].width = 16

    wb.save(os.path.join(OUTPUT_DIR, "11_Property_Comparison_Sheet.xlsx"))
    print("Created: 11_Property_Comparison_Sheet.xlsx")


# ──────────────────────────────────────────────
# XLSX 2: Price Trends Data
# ──────────────────────────────────────────────
def create_price_trends_xlsx():
    wb = Workbook()
    ws = wb.active
    ws.title = "Price Trends"

    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="0D6CF2", end_color="0D6CF2", fill_type="solid")

    headers = ["Area", "2020 (Rs/sqft)", "2021 (Rs/sqft)", "2022 (Rs/sqft)",
               "2023 (Rs/sqft)", "2024 (Rs/sqft)", "5yr Growth %", "Avg Rental (Rs/sqft/month)"]

    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center", wrap_text=True)

    areas = [
        ["Whitefield", 4200, 4600, 5500, 6800, 7800, "86%", 32],
        ["Electronic City", 3200, 3500, 4000, 4800, 5600, "75%", 25],
        ["Sarjapur Road", 3800, 4200, 5000, 5800, 6800, "79%", 28],
        ["Hebbal", 5500, 6000, 6800, 7800, 8800, "60%", 38],
        ["Koramangala", 9500, 10200, 11500, 13000, 14500, "53%", 55],
        ["HSR Layout", 6200, 6800, 7500, 8500, 9800, "58%", 42],
        ["Bellandur", 5000, 5600, 6500, 7800, 9000, "80%", 38],
        ["Yelahanka", 3500, 3800, 4300, 5000, 5800, "66%", 22],
        ["Kanakapura Road", 2800, 3100, 3500, 4000, 4600, "64%", 18],
        ["Devanahalli", 2500, 2800, 3200, 3800, 4500, "80%", 16],
        ["Marathahalli", 5200, 5700, 6500, 7500, 8500, "63%", 35],
        ["Indiranagar", 11000, 12000, 13500, 15000, 16500, "50%", 60],
        ["Jayanagar", 8000, 8700, 9500, 10800, 12000, "50%", 45],
        ["Bannerghatta Road", 3800, 4200, 4800, 5500, 6500, "71%", 26],
    ]

    for row_idx, row_data in enumerate(areas, 2):
        for col_idx, value in enumerate(row_data, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=value)
            cell.alignment = Alignment(horizontal="center")

    for col in range(1, len(headers) + 1):
        ws.column_dimensions[ws.cell(row=1, column=col).column_letter].width = 18

    wb.save(os.path.join(OUTPUT_DIR, "12_Price_Trends_Bangalore.xlsx"))
    print("Created: 12_Price_Trends_Bangalore.xlsx")


# ──────────────────────────────────────────────
# DOCX 1: Property Buying Guide
# ──────────────────────────────────────────────
def create_buying_guide_docx():
    doc = Document()

    title = doc.add_heading("Complete Property Buying Guide - Bangalore", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("1. Understanding Property Types", level=1)
    doc.add_paragraph(
        "Bangalore's real estate market offers diverse property types to suit different buyer needs and budgets. "
        "Understanding these types is crucial before making a purchase decision."
    )

    doc.add_heading("Apartments", level=2)
    doc.add_paragraph(
        "Apartments are the most common residential property type in Bangalore, ranging from affordable 1BHK units "
        "starting at Rs. 25-30 Lakhs in peripheral areas to luxury penthouses costing Rs. 5+ Crores in prime locations. "
        "Apartments offer shared amenities, security, and community living. Typical maintenance charges range from "
        "Rs. 2-5 per sq.ft. per month. Key considerations include floor level (higher floors command 2-5% premium), "
        "facing direction (East and North facing are preferred), and parking allocation."
    )

    doc.add_heading("Independent Villas", level=2)
    doc.add_paragraph(
        "Villas provide maximum privacy and space, typically ranging from Rs. 1.5-6 Crores in Bangalore. "
        "They come with private gardens, dedicated parking, and independent utility connections. Popular villa "
        "clusters are in Sarjapur Road, Whitefield, and Yelahanka. Maintenance is the owner's responsibility, "
        "though gated villa communities share common area maintenance costs."
    )

    doc.add_heading("Plotted Developments", level=2)
    doc.add_paragraph(
        "Buying a plot gives maximum flexibility to build a custom home. BMRDA/BDA-approved plots in Bangalore "
        "range from Rs. 3,000-8,000 per sq.ft. depending on location. Ensure the plot has clear title, approved "
        "layout plan, and all utility connections. Popular areas for plots include Kanakapura Road, Devanahalli, "
        "and Sarjapur. Bank loans are available for approved layouts."
    )

    doc.add_heading("2. Legal Checklist", level=1)
    items = [
        "RERA Registration: Every project must be registered with Karnataka RERA (K-RERA). Verify at rera.karnataka.gov.in",
        "Title Deed: Ensure clear ownership chain for at least 30 years. Get title verified by an independent lawyer.",
        "Encumbrance Certificate (EC): Obtain for the last 30 years to confirm no pending loans or disputes.",
        "Khata Certificate: Essential for property tax payment and obtaining building plan approval.",
        "Completion Certificate (CC) & Occupancy Certificate (OC): Must be obtained before possession.",
        "Building Plan Approval: Verify from BBMP/BDA/BMRDA depending on the jurisdiction.",
        "Land Use Certificate: Confirm the land is zoned for residential/commercial use as applicable.",
        "NOC from various departments: Fire, water, electricity, environment clearance.",
    ]
    for item in items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Home Loan Guide", level=1)
    doc.add_paragraph(
        "Most banks offer home loans up to 80% of property value (75% for properties above Rs. 75 Lakhs). "
        "Current interest rates range from 8.35-9.5% per annum. Key banks and their current rates:"
    )

    table = doc.add_table(rows=8, cols=4)
    table.style = "Table Grid"
    headers = ["Bank", "Interest Rate", "Max Tenure", "Processing Fee"]
    for j, header in enumerate(headers):
        table.rows[0].cells[j].text = header
        table.rows[0].cells[j].paragraphs[0].runs[0].font.bold = True if table.rows[0].cells[j].paragraphs[0].runs else None

    bank_data = [
        ["SBI", "8.35% onwards", "30 years", "0.35% of loan amount"],
        ["HDFC Bank", "8.50% onwards", "30 years", "0.50% of loan amount"],
        ["ICICI Bank", "8.45% onwards", "30 years", "0.50% of loan amount"],
        ["Axis Bank", "8.55% onwards", "25 years", "Rs. 10,000 flat"],
        ["Bank of Baroda", "8.40% onwards", "30 years", "0.25% of loan amount"],
        ["Kotak Mahindra", "8.60% onwards", "25 years", "0.50% of loan amount"],
        ["LIC Housing", "8.50% onwards", "30 years", "Rs. 10,000-15,000"],
    ]
    for i, row_data in enumerate(bank_data, 1):
        for j, val in enumerate(row_data):
            table.rows[i].cells[j].text = val

    doc.add_heading("4. Tax Benefits", level=1)
    doc.add_paragraph(
        "Home buyers can avail significant tax deductions under the Income Tax Act:\n\n"
        "Section 80C: Principal repayment - up to Rs. 1.5 Lakhs per year\n"
        "Section 24(b): Interest on home loan - up to Rs. 2 Lakhs per year (self-occupied property)\n"
        "Section 80EEA: Additional Rs. 1.5 Lakhs interest deduction for first-time buyers (stamp duty value up to Rs. 45 Lakhs)\n\n"
        "Registration charges in Karnataka: 5.6% of property value (5% stamp duty + 1% registration + surcharge). "
        "GST of 5% applies to under-construction properties (1% for affordable housing under Rs. 45 Lakhs)."
    )

    doc.save(os.path.join(OUTPUT_DIR, "13_Property_Buying_Guide.docx"))
    print("Created: 13_Property_Buying_Guide.docx")


# ──────────────────────────────────────────────
# DOCX 2: Amenities Specification Document
# ──────────────────────────────────────────────
def create_amenities_spec_docx():
    doc = Document()

    title = doc.add_heading("Standard Amenities Specification Guide", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Reference Document for Premium Residential Projects in Bangalore").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("Swimming Pool Standards", level=1)
    doc.add_paragraph(
        "Main Pool: Minimum 25m x 10m lap pool with depth ranging from 3.5 ft to 5.5 ft. Pool must have "
        "6-lane markings for competitive swimming. Water filtration using sand filter and UV treatment. "
        "Temperature control system maintaining 26-28 degrees Celsius. Pool deck with Italian porcelain "
        "anti-skid tiles. Minimum 12 sun loungers and 4 umbrellas.\n\n"
        "Kids Pool: Separate pool with maximum depth of 2 feet. Mushroom shower and water play features. "
        "Soft rubber flooring around the pool area. Dedicated lifeguard during operating hours (6 AM - 9 PM).\n\n"
        "Changing Rooms: Separate male and female changing rooms with lockers, showers, and washrooms. "
        "Minimum 20 lockers per changing room."
    )

    doc.add_heading("Gymnasium Standards", level=1)
    doc.add_paragraph(
        "Minimum area: 3,000 sq.ft. for projects with 300+ units. Equipment must be commercial grade "
        "(Technogym, Life Fitness, or equivalent). Required equipment includes:\n\n"
        "Cardio Zone: 8 treadmills, 6 ellipticals, 4 stationary bikes, 2 rowing machines, 2 stair climbers. "
        "All cardio equipment must have individual entertainment screens.\n\n"
        "Strength Zone: Complete dumbbell rack (5-50 kg), Olympic barbell sets, Smith machine, cable crossover, "
        "leg press, lat pulldown, seated row, chest press, shoulder press machines. Minimum 2 adjustable benches.\n\n"
        "Free Weights Area: Rubber matted floor, mirror wall, kettlebell set (4-32 kg).\n\n"
        "Functional Training Zone: TRX suspension trainers, battle ropes, medicine balls, resistance bands, "
        "plyo boxes.\n\n"
        "The gym must have professional-grade flooring (rubber tiles, 8mm minimum thickness), "
        "air conditioning maintaining 22-24 degrees, water dispenser, and certified trainer on duty."
    )

    doc.add_heading("Clubhouse Standards", level=1)
    doc.add_paragraph(
        "Minimum clubhouse area: 5,000 sq.ft. for projects with 300+ units. The clubhouse must include:\n\n"
        "Banquet Hall: Capacity for minimum 150 guests. Air-conditioned with stage area. Professional sound "
        "system and projection facility. Catering kitchen adjacent to the hall.\n\n"
        "Indoor Games Room: Billiards table (tournament size), 2 table tennis tables, carrom boards, chess "
        "tables, foosball table. Air-conditioned with appropriate lighting.\n\n"
        "Library/Reading Lounge: Minimum 1,000 sq.ft. with seating for 20 people. Wi-Fi enabled. "
        "Subscription to 10+ magazines and 5+ newspapers.\n\n"
        "Multi-purpose Room: 800 sq.ft. flexible space for yoga, dance, art classes. Mirror wall, "
        "wooden flooring, and sound system."
    )

    doc.add_heading("Security Infrastructure", level=1)
    doc.add_paragraph(
        "Entry/Exit Control: Boom barriers with RFID/ANPR at all vehicle entry points. Pedestrian turnstiles "
        "with biometric/card access. Visitor management system with photo capture and OTP verification.\n\n"
        "CCTV Coverage: HD cameras at all entry/exit points, parking levels, elevator lobbies, common areas, "
        "and perimeter. Minimum 30-day recording storage. Night vision capability. AI-powered analytics for "
        "intrusion detection and abandoned object detection.\n\n"
        "Intercom System: IP-based intercom connecting every apartment to the security room, clubhouse, "
        "and management office. Video door phone at apartment entrance connected to guard room.\n\n"
        "Fire Safety: Smoke detectors in every apartment. Sprinkler system in common areas and parking. "
        "Fire extinguishers on every floor. Fire alarm system connected to nearest fire station. "
        "Two dedicated fire exits per floor. Pressurized stairwells. Fire pump room with diesel and electric pumps."
    )

    doc.add_heading("Landscaping Standards", level=1)
    doc.add_paragraph(
        "Minimum 65% of total project area must be open space (excluding building footprint and roads). "
        "Landscaping must include a mix of native trees (minimum 50% of tree species must be native), "
        "ornamental plants, and lawned areas.\n\n"
        "Required green features: Central garden/park (minimum 1 acre for projects above 5 acres), "
        "jogging/walking track (rubberized surface, minimum 500m), children's play area (imported equipment, "
        "EPDM rubber flooring), senior citizen area with reflexology path, pet-friendly zone with agility "
        "equipment, organic composting facility, and drip irrigation system for all landscaped areas."
    )

    doc.save(os.path.join(OUTPUT_DIR, "14_Amenities_Specification_Guide.docx"))
    print("Created: 14_Amenities_Specification_Guide.docx")


# ──────────────────────────────────────────────
# Run all generators
# ──────────────────────────────────────────────
if __name__ == "__main__":
    print("Generating real estate documents...\n")

    create_sunrise_heights()
    create_green_valley()
    create_metro_edge()
    create_royal_orchid()
    create_eco_habitat()
    create_lakeshore()
    create_heritage_grand()
    create_skyline_commercial()
    create_palm_grove()
    create_market_report()
    create_property_comparison_xlsx()
    create_price_trends_xlsx()
    create_buying_guide_docx()
    create_amenities_spec_docx()

    print(f"\nDone! All documents saved to: {os.path.abspath(OUTPUT_DIR)}")
    print(f"Total files: {len(os.listdir(OUTPUT_DIR))}")
