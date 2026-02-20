"""
Document parsers for streaming PDF, DOCX, and Excel files.
No disk storage - all processing done in-memory from bytes.
"""

from typing import Union
import io


# =============================================================================
# PDF PARSERS
# =============================================================================


def parse_pdf_pypdf2(file_bytes: bytes) -> str:
    """
    Parse PDF file using PyPDF2 (simple, reliable).

    Args:
        file_bytes: PDF file as bytes

    Returns:
        Extracted text from all pages
    """
    try:
        from PyPDF2 import PdfReader
    except ImportError:
        raise ImportError("PyPDF2 not installed. Run: pip install PyPDF2")

    # Create file-like object from bytes
    pdf_file = io.BytesIO(file_bytes)

    # Read PDF
    reader = PdfReader(pdf_file)
    text_parts = []

    # Extract text from all pages
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text.strip():
            text_parts.append(f"[Page {page_num}]\n{text}")

    return "\n\n".join(text_parts)


def parse_pdf_pdfplumber(file_bytes: bytes) -> str:
    """
    Parse PDF file using pdfplumber (better for tables, complex layouts).

    Args:
        file_bytes: PDF file as bytes

    Returns:
        Extracted text from all pages
    """
    try:
        import pdfplumber
    except ImportError:
        raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")

    # Create file-like object from bytes
    pdf_file = io.BytesIO(file_bytes)
    text_parts = []

    # Read PDF
    with pdfplumber.open(pdf_file) as pdf:
        for page_num, page in enumerate(pdf.pages, 1):
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(f"[Page {page_num}]\n{text}")

            # Extract tables if present
            tables = page.extract_tables()
            if tables:
                for table_num, table in enumerate(tables, 1):
                    table_text = "\n".join(["\t".join(str(cell) for cell in row) for row in table])
                    text_parts.append(f"[Page {page_num} - Table {table_num}]\n{table_text}")

    return "\n\n".join(text_parts)


def parse_pdf_stream(file_bytes: bytes, parser: str = "pypdf2") -> str:
    """
    Parse PDF from bytes stream.

    Args:
        file_bytes: PDF file as bytes
        parser: Parser to use ("pypdf2" or "pdfplumber")

    Returns:
        Extracted text
    """
    if parser == "pdfplumber":
        return parse_pdf_pdfplumber(file_bytes)
    else:
        return parse_pdf_pypdf2(file_bytes)


# =============================================================================
# DOCX PARSER
# =============================================================================


def parse_docx_stream(file_bytes: bytes) -> str:
    """
    Parse DOCX file from bytes stream.

    Args:
        file_bytes: DOCX file as bytes

    Returns:
        Extracted text from all paragraphs
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed. Run: pip install python-docx")

    # Create file-like object from bytes
    docx_file = io.BytesIO(file_bytes)

    # Read DOCX
    doc = Document(docx_file)
    text_parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text_parts.append(para.text)

    # Extract tables
    for table_num, table in enumerate(doc.tables, 1):
        table_text = []
        for row in table.rows:
            row_text = "\t".join(cell.text for cell in row.cells)
            table_text.append(row_text)
        if table_text:
            text_parts.append(f"\n[Table {table_num}]\n" + "\n".join(table_text))

    return "\n\n".join(text_parts)


# =============================================================================
# EXCEL PARSERS
# =============================================================================


def parse_excel_openpyxl(file_bytes: bytes, combine_sheets: bool = True) -> str:
    """
    Parse Excel file using openpyxl (supports .xlsx).

    Args:
        file_bytes: Excel file as bytes
        combine_sheets: If True, combine all sheets; if False, separate by sheet

    Returns:
        Extracted text from all sheets
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("openpyxl not installed. Run: pip install openpyxl")

    # Create file-like object from bytes
    excel_file = io.BytesIO(file_bytes)

    # Load workbook
    wb = load_workbook(excel_file, read_only=True, data_only=True)
    text_parts = []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        sheet_text = []

        # Extract all rows
        for row in sheet.iter_rows(values_only=True):
            # Filter out None and empty values
            row_values = [str(cell) for cell in row if cell is not None and str(cell).strip()]
            if row_values:
                sheet_text.append("\t".join(row_values))

        if sheet_text:
            if combine_sheets:
                text_parts.extend(sheet_text)
            else:
                text_parts.append(f"[Sheet: {sheet_name}]\n" + "\n".join(sheet_text))

    return "\n\n".join(text_parts)


def parse_excel_pandas(file_bytes: bytes, combine_sheets: bool = True) -> str:
    """
    Parse Excel file using pandas (simpler API, good for data).

    Args:
        file_bytes: Excel file as bytes
        combine_sheets: If True, combine all sheets; if False, separate by sheet

    Returns:
        Extracted text from all sheets
    """
    try:
        import pandas as pd
    except ImportError:
        raise ImportError("pandas not installed. Run: pip install pandas openpyxl")

    # Create file-like object from bytes
    excel_file = io.BytesIO(file_bytes)

    # Read all sheets
    excel_data = pd.read_excel(excel_file, sheet_name=None, engine='openpyxl')
    text_parts = []

    for sheet_name, df in excel_data.items():
        # Convert dataframe to text
        sheet_text = df.to_string(index=False, na_rep='')

        if combine_sheets:
            text_parts.append(sheet_text)
        else:
            text_parts.append(f"[Sheet: {sheet_name}]\n{sheet_text}")

    return "\n\n".join(text_parts)


def parse_excel_stream(
    file_bytes: bytes,
    parser: str = "pandas",
    combine_sheets: bool = True
) -> str:
    """
    Parse Excel from bytes stream.

    Args:
        file_bytes: Excel file as bytes
        parser: Parser to use ("pandas" or "openpyxl")
        combine_sheets: If True, combine all sheets

    Returns:
        Extracted text
    """
    if parser == "openpyxl":
        return parse_excel_openpyxl(file_bytes, combine_sheets)
    else:
        return parse_excel_pandas(file_bytes, combine_sheets)


# =============================================================================
# TEXT FILE PARSER
# =============================================================================


def parse_txt_stream(file_bytes: bytes) -> str:
    """
    Parse plain text file from bytes.

    Args:
        file_bytes: Text file as bytes

    Returns:
        Decoded text
    """
    # Try UTF-8 first, fallback to latin-1
    try:
        return file_bytes.decode('utf-8')
    except UnicodeDecodeError:
        return file_bytes.decode('latin-1')


# =============================================================================
# AUTO-DETECTION AND ROUTING
# =============================================================================


def auto_detect_and_parse(
    file_bytes: bytes,
    filename: str,
    pdf_parser: str = "pypdf2",
    excel_parser: str = "pandas",
    excel_combine_sheets: bool = True
) -> str:
    """
    Auto-detect file type and parse accordingly.

    Args:
        file_bytes: File content as bytes
        filename: Original filename (used to detect type)
        pdf_parser: PDF parser to use ("pypdf2" or "pdfplumber")
        excel_parser: Excel parser to use ("pandas" or "openpyxl")
        excel_combine_sheets: Combine Excel sheets or separate

    Returns:
        Extracted text

    Raises:
        ValueError: If file type is not supported
    """
    # Get file extension
    filename_lower = filename.lower()

    if filename_lower.endswith('.pdf'):
        return parse_pdf_stream(file_bytes, parser=pdf_parser)

    elif filename_lower.endswith('.docx'):
        return parse_docx_stream(file_bytes)

    elif filename_lower.endswith(('.xlsx', '.xls')):
        return parse_excel_stream(
            file_bytes,
            parser=excel_parser,
            combine_sheets=excel_combine_sheets
        )

    elif filename_lower.endswith('.txt'):
        return parse_txt_stream(file_bytes)

    else:
        raise ValueError(
            f"Unsupported file type: {filename}. "
            f"Supported: .pdf, .docx, .xlsx, .xls, .txt"
        )


# =============================================================================
# UTILITY FUNCTIONS
# =============================================================================


def validate_file_size(file_bytes: bytes, max_size_mb: int = 50) -> bool:
    """
    Validate file size.

    Args:
        file_bytes: File content as bytes
        max_size_mb: Maximum allowed size in MB

    Returns:
        True if valid, False otherwise
    """
    size_mb = len(file_bytes) / (1024 * 1024)
    return size_mb <= max_size_mb


def get_file_info(file_bytes: bytes, filename: str) -> dict:
    """
    Get file information.

    Args:
        file_bytes: File content as bytes
        filename: Original filename

    Returns:
        Dictionary with file info
    """
    size_mb = len(file_bytes) / (1024 * 1024)
    extension = filename.split('.')[-1].lower() if '.' in filename else 'unknown'

    return {
        "filename": filename,
        "extension": extension,
        "size_bytes": len(file_bytes),
        "size_mb": round(size_mb, 2),
    }


# =============================================================================
# TEXT CLEANING
# =============================================================================


def clean_text(text: str) -> str:
    """
    Clean extracted text before chunking.
    Removes PDF artifacts, normalizes whitespace, fixes hyphenation.

    Args:
        text: Raw extracted text

    Returns:
        Cleaned text
    """
    import re

    # Remove page markers injected by parsers (e.g. "[Page 1]")
    text = re.sub(r'\[Page \d+\]\n?', '', text)

    # Remove table markers (e.g. "[Table 1]", "[Page 2 - Table 1]")
    text = re.sub(r'\[(?:Page \d+ - )?Table \d+\]\n?', '', text)

    # Fix hyphenation across line breaks (e.g. "apart-\nment" -> "apartment")
    text = re.sub(r'(\w)-\n(\w)', r'\1\2', text)

    # Collapse multiple blank lines into one
    text = re.sub(r'\n{3,}', '\n\n', text)

    # Collapse multiple spaces/tabs into single space
    text = re.sub(r'[ \t]{2,}', ' ', text)

    # Remove leading/trailing whitespace on each line
    text = '\n'.join(line.strip() for line in text.split('\n'))

    return text.strip()


# =============================================================================
# TEXT CHUNKING
# =============================================================================


def chunk_text(text: str, chunk_size: int = 800, overlap: int = 200) -> list:
    """
    Split text into overlapping chunks at sentence boundaries.
    Respects sentence endings so chunks never cut mid-sentence.

    Args:
        text: Text to chunk
        chunk_size: Target size of each chunk in characters
        overlap: Target overlap between consecutive chunks in characters

    Returns:
        List of text chunks
    """
    import re

    # Split into sentences (handles ., !, ?, and common abbreviations)
    sentences = re.split(r'(?<=[.!?])\s+', text)
    sentences = [s.strip() for s in sentences if s.strip()]

    if not sentences:
        return []

    chunks = []
    current_sentences = []
    current_length = 0

    for sentence in sentences:
        # If adding this sentence exceeds chunk_size and we already have content,
        # finalize the current chunk
        if current_length + len(sentence) + 1 > chunk_size and current_sentences:
            chunks.append(' '.join(current_sentences))

            # Build overlap: keep trailing sentences that fit within the overlap budget
            overlap_sentences = []
            overlap_length = 0
            for s in reversed(current_sentences):
                if overlap_length + len(s) + 1 > overlap:
                    break
                overlap_sentences.insert(0, s)
                overlap_length += len(s) + 1

            current_sentences = overlap_sentences
            current_length = overlap_length

        current_sentences.append(sentence)
        current_length += len(sentence) + 1

    # Don't forget the last chunk
    if current_sentences:
        chunks.append(' '.join(current_sentences))

    return [c for c in chunks if c.strip()]


if __name__ == "__main__":
    # Test with sample data
    print("Document Parsers Test")
    print("=" * 80)

    # Test text chunking
    sample_text = "This is a sample text for testing chunking functionality. " * 20
    chunks = chunk_text(sample_text, chunk_size=100, overlap=20)
    print(f"\nChunking Test:")
    print(f"Original text length: {len(sample_text)} chars")
    print(f"Number of chunks: {len(chunks)}")
    print(f"First chunk: {chunks[0][:50]}...")

    # Test file info
    test_bytes = b"Sample file content"
    info = get_file_info(test_bytes, "test.pdf")
    print(f"\nFile Info Test:")
    print(info)
